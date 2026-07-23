#!/usr/bin/env python3
"""Build the evidence-backed Song cross-code comparison Word report."""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
SUMMARY_PATH = REPO_ROOT / ".tmp" / "final-results" / "comparison_summary.json"
OVERVIEW_PATH = REPO_ROOT / ".tmp" / "final-results" / "cross_code_comparison.png"
CASE_PATH = REPO_ROOT / "benchmarks" / "sotem" / "song2025_layered_pair.yaml"
OUTPUT_PATH = (
    REPO_ROOT
    / "output"
    / "doc"
    / "Song文献模型_FEniCSx_empymod_SimPEG_三维时域正演对比报告.docx"
)
WORK_DIR = REPO_ROOT / "tmp" / "docs" / "song_cross_code_report"
SONG_FIGURE_DIR = REPO_ROOT / "assets" / "reports" / "song2025"
SONG_PAPER_FIGURES = {
    "layout": SONG_FIGURE_DIR / "song2025_fig1_layout.png",
    "model": SONG_FIGURE_DIR / "song2025_fig4_model.png",
    "ex": SONG_FIGURE_DIR / "song2025_fig7_ex.png",
    "hz": SONG_FIGURE_DIR / "song2025_fig8_hz.png",
}

WSL_ROOT = Path(r"\\wsl.localhost\Ubuntu\home\paidaxin")
FENICSX_DIRS = {
    "noip": WSL_ROOT
    / "codex-sotem-song-be-noip-mpi12-t8-1ms-5d03d6e"
    / "song-noip-be-t21-1ms",
    "ip": WSL_ROOT
    / "codex-sotem-song-be-ip-d8-mpi12-t8-1ms-5d03d6e"
    / "song-ip-be-t21-1ms",
}
SIMPEG_DIRS = {
    "noip": WSL_ROOT / "codex-song-simpeg-1ms-noip-s1t16-5d03d6e",
    "ip": WSL_ROOT / "codex-song-simpeg-1ms-ip-s1t16-5d03d6e",
}

COMPONENTS = ("Ex", "Hz", "dBzdt")
COMPONENT_LABELS = {
    "Ex": r"$|E_x|$ (V/m)",
    "Hz": r"$|H_z|$ (A/m)",
    "dBzdt": r"$|dB_z/dt|$ (T/s)",
}
COMPONENT_CN = {"Ex": "Ex", "Hz": "Hz", "dBzdt": "dBz/dt"}
GATE = 0.05

EXPECTED_MAX = {
    "fenicsx_noip": {
        "Ex": 0.01036429269070334,
        "Hz": 0.014124974665132994,
        "dBzdt": 0.03301669355706037,
    },
    "fenicsx_ip": {
        "Ex": 0.01993756177376695,
        "Hz": 0.010906898025984593,
        "dBzdt": 0.02687928048666212,
    },
    "simpeg_noip": {
        "Ex": 0.032101056843179654,
        "Hz": 0.006303089535895287,
        "dBzdt": 0.034527646600017094,
    },
    "simpeg_ip": {
        "Ex": 0.07420634839988084,
        "Hz": 0.006542826240444855,
        "dBzdt": 0.034454769506365814,
    },
}

FIGURE_CONTRACT = {
    "core_conclusion": (
        "FEniCSx reproduces the finite-wire reference below the fixed 5% gate "
        "for all formal no-IP and IP components; only SimPEG IP Ex fails at 10 us."
    ),
    "archetype": "quantitative grid with a supporting model schematic",
    "backend": "Python/matplotlib",
    "final_width_mm": 170,
    "source_data": "all 21 observations from 10 us to 1 ms",
    "integrity": "no shifting, scaling, smoothing, point deletion, or gate changes",
    "reviewer_risk": (
        "empymod reference role, 1 ms scope, point receiver, ideal step-off, "
        "and retained SimPEG IP Ex failure"
    ),
}


@dataclass(frozen=True)
class VariantData:
    times: np.ndarray
    solver: np.ndarray
    reference: np.ndarray
    components: tuple[str, ...]

    def column(self, values: np.ndarray, component: str) -> np.ndarray:
        return values[:, self.components.index(component)]


@dataclass
class Evidence:
    summary: dict[str, Any]
    case: dict[str, Any]
    fenicsx: dict[str, VariantData]
    simpeg: dict[str, VariantData]
    fenicsx_config: dict[str, dict[str, Any]]
    fenicsx_errors: dict[str, dict[str, Any]]
    fenicsx_diagnostics: dict[str, dict[str, Any]]
    fenicsx_mesh: dict[str, dict[str, Any]]
    fenicsx_magnetic: dict[str, dict[str, Any]]
    simpeg_json: dict[str, dict[str, Any]]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def linux_to_unc(path_string: str) -> Path:
    prefix = "/home/paidaxin/"
    if not path_string.startswith(prefix):
        raise ValueError(f"Unsupported evidence path: {path_string}")
    return WSL_ROOT / path_string[len(prefix) :]


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def load_yaml(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def load_fenicsx_npz(path: Path) -> VariantData:
    with np.load(path) as payload:
        return VariantData(
            times=np.asarray(payload["times"], dtype=float),
            solver=np.asarray(payload["fem"], dtype=float),
            reference=np.asarray(payload["empymod"], dtype=float),
            components=tuple(str(value) for value in payload["components"].tolist()),
        )


def load_simpeg_npz(path: Path) -> VariantData:
    with np.load(path) as payload:
        return VariantData(
            times=np.asarray(payload["times"], dtype=float),
            solver=np.asarray(payload["simpeg"], dtype=float),
            reference=np.asarray(payload["reference"], dtype=float),
            components=tuple(str(value) for value in payload["components"].tolist()),
        )


def verify_variant(variant: VariantData, label: str) -> None:
    if variant.times.shape != (21,):
        raise AssertionError(f"{label}: expected 21 observations, got {variant.times.shape}")
    np.testing.assert_allclose(variant.times[0], 1.0e-5, rtol=0, atol=1.0e-15)
    np.testing.assert_allclose(variant.times[-1], 1.0e-3, rtol=0, atol=1.0e-15)
    if variant.solver.shape != (21, 4) or variant.reference.shape != (21, 4):
        raise AssertionError(f"{label}: response arrays must be 21 x 4")
    for component in ("Ex", "Ey", "Hz", "dBzdt"):
        if component not in variant.components:
            raise AssertionError(f"{label}: missing {component}")
    if not np.isfinite(variant.solver).all() or not np.isfinite(variant.reference).all():
        raise AssertionError(f"{label}: non-finite response value")


def load_evidence() -> Evidence:
    required = [SUMMARY_PATH, OVERVIEW_PATH, CASE_PATH]
    required.extend(SONG_PAPER_FIGURES.values())
    for directory in [*FENICSX_DIRS.values(), *SIMPEG_DIRS.values()]:
        required.append(directory)
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing evidence:\n" + "\n".join(missing))

    summary = load_json(SUMMARY_PATH)
    if summary.get("schema") != "song_1ms_cross_code_comparison/v1":
        raise AssertionError("Unexpected comparison summary schema")
    if summary.get("sample_count") != 21:
        raise AssertionError("Comparison summary must contain 21 observations")
    np.testing.assert_allclose(summary["time_window_s"], [1.0e-5, 1.0e-3], rtol=0, atol=1e-15)
    if float(summary.get("relative_error_gate")) != GATE:
        raise AssertionError("The formal error gate changed")

    for variant_name, components in EXPECTED_MAX.items():
        for component, expected in components.items():
            actual = summary["comparisons"][variant_name][component]["max_relative_error"]
            np.testing.assert_allclose(actual, expected, rtol=0, atol=5.0e-12)

    fenicsx: dict[str, VariantData] = {}
    simpeg: dict[str, VariantData] = {}
    for variant in ("noip", "ip"):
        f_path = FENICSX_DIRS[variant] / "verification_data.npz"
        s_path = SIMPEG_DIRS[variant] / "result.npz"
        expected_f_hash = summary["inputs"][f"fenicsx_{variant}"]["sha256"]
        expected_s_hash = summary["inputs"][f"simpeg_{variant}"]["sha256"]
        if sha256_file(f_path) != expected_f_hash:
            raise AssertionError(f"FEniCSx {variant} hash mismatch")
        if sha256_file(s_path) != expected_s_hash:
            raise AssertionError(f"SimPEG {variant} hash mismatch")
        fenicsx[variant] = load_fenicsx_npz(f_path)
        simpeg[variant] = load_simpeg_npz(s_path)
        verify_variant(fenicsx[variant], f"FEniCSx {variant}")
        verify_variant(simpeg[variant], f"SimPEG {variant}")
        np.testing.assert_allclose(fenicsx[variant].times, simpeg[variant].times, rtol=0, atol=1e-15)

    return Evidence(
        summary=summary,
        case=load_yaml(CASE_PATH),
        fenicsx=fenicsx,
        simpeg=simpeg,
        fenicsx_config={
            key: load_yaml(directory / "run_config_resolved.yaml")
            for key, directory in FENICSX_DIRS.items()
        },
        fenicsx_errors={
            key: load_json(directory / "error_summary.json")
            for key, directory in FENICSX_DIRS.items()
        },
        fenicsx_diagnostics={
            key: load_json(directory / "diagnostics.json")
            for key, directory in FENICSX_DIRS.items()
        },
        fenicsx_mesh={
            key: load_json(directory / "mesh_quality_preflight.json")
            for key, directory in FENICSX_DIRS.items()
        },
        fenicsx_magnetic={
            key: load_json(directory / "magnetic_initialization.json")
            for key, directory in FENICSX_DIRS.items()
        },
        simpeg_json={
            key: load_json(directory / "result.json")
            for key, directory in SIMPEG_DIRS.items()
        },
    )


def configure_matplotlib() -> None:
    mpl.rcParams.update(
        {
            "font.family": "sans-serif",
            "font.sans-serif": [
                "Microsoft YaHei",
                "SimHei",
                "Arial",
                "DejaVu Sans",
                "sans-serif",
            ],
            "axes.unicode_minus": False,
            "svg.fonttype": "none",
            "pdf.fonttype": 42,
            "font.size": 9,
            "axes.labelsize": 9,
            "axes.titlesize": 10,
            "xtick.labelsize": 8,
            "ytick.labelsize": 8,
            "legend.fontsize": 8,
            "axes.spines.right": False,
            "axes.spines.top": False,
            "axes.linewidth": 0.8,
            "legend.frameon": False,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "savefig.facecolor": "white",
        }
    )


def save_figure(fig: plt.Figure, stem: Path) -> Path:
    stem.parent.mkdir(parents=True, exist_ok=True)
    png = stem.with_suffix(".png")
    fig.savefig(png, dpi=300, bbox_inches="tight")
    fig.savefig(stem.with_suffix(".svg"), bbox_inches="tight")
    fig.savefig(stem.with_suffix(".pdf"), bbox_inches="tight")
    plt.close(fig)
    if png.stat().st_size < 20_000:
        raise AssertionError(f"Figure output unexpectedly small: {png}")
    return png


def component_values(data: VariantData, source: str, component: str) -> np.ndarray:
    array = data.solver if source == "solver" else data.reference
    return data.column(array, component)


def relative_error(values: np.ndarray, reference: np.ndarray, floor: float) -> np.ndarray:
    return np.abs(values - reference) / np.maximum(np.abs(reference), floor)


def draw_model_schematic(evidence: Evidence, output_dir: Path) -> Path:
    fig, (ax_top, ax_depth) = plt.subplots(1, 2, figsize=(7.2, 3.4), gridspec_kw={"width_ratios": [1.45, 1]})

    ax_top.axhline(0, color="#9aa0a6", lw=0.8, zorder=0)
    ax_top.plot([-500, 500], [0, 0], color="#2b6cb0", lw=5, solid_capstyle="round")
    ax_top.scatter([-500, 500], [0, 0], s=90, color="#2b6cb0", edgecolor="white", zorder=3)
    ax_top.scatter([0], [-500], marker="^", s=110, color="#d97706", edgecolor="white", zorder=4)
    ax_top.annotate("A (-500, 0)", (-500, 0), xytext=(-530, 90), ha="center", arrowprops={"arrowstyle": "-", "lw": 0.7})
    ax_top.annotate("B (500, 0)", (500, 0), xytext=(530, 90), ha="center", arrowprops={"arrowstyle": "-", "lw": 0.7})
    ax_top.annotate("Rx (0, -500)", (0, -500), xytext=(180, -610), ha="center", arrowprops={"arrowstyle": "->", "lw": 0.8})
    ax_top.annotate("1000 m有限接地线源，10 A理想step-off", (0, 0), xytext=(0, 180), ha="center", color="#1e4f8a")
    ax_top.annotate("500 m", (0, -250), xytext=(70, -260), color="#555555")
    ax_top.annotate("", xy=(0, -495), xytext=(0, -5), arrowprops={"arrowstyle": "<->", "lw": 1.0, "color": "#555555"})
    ax_top.set_xlim(-700, 700)
    ax_top.set_ylim(-700, 260)
    ax_top.set_aspect("equal", adjustable="box")
    ax_top.set_xlabel("x (m)")
    ax_top.set_ylabel("y (m)")
    ax_top.set_title("a  平面几何")
    ax_top.grid(alpha=0.18)

    ax_depth.axhspan(-250, 0, color="#edf6ff", label="空气")
    ax_depth.axhspan(0, 300, color="#f7dfbd", label="0-300 m参数区")
    ax_depth.axhspan(300, 700, color="#ead0a8", label="下伏半空间")
    ax_depth.axhline(0, color="#444444", lw=1.2)
    ax_depth.axhline(300, color="#8a5a1f", lw=1.0, ls="--")
    ax_depth.text(0.5, -110, "空气：ρ = 10^6 Ω·m", ha="center", va="center")
    ax_depth.text(0.5, 135, "ρ = 100 Ω·m\nIP：m = 0.3，τ = 1 s，c = 0.3", ha="center", va="center")
    ax_depth.text(0.5, 485, "ρ = 100 Ω·m\n无极化", ha="center", va="center")
    ax_depth.annotate("z向下", xy=(0.92, 600), xytext=(0.92, 380), ha="center", arrowprops={"arrowstyle": "->", "lw": 1.1})
    ax_depth.set_xlim(0, 1)
    ax_depth.set_ylim(700, -250)
    ax_depth.set_xticks([])
    ax_depth.set_ylabel("深度 z (m)")
    ax_depth.set_title("b  分层与极化参数")
    for spine in ax_depth.spines.values():
        spine.set_visible(True)
        spine.set_linewidth(0.8)

    fig.suptitle("Song文献模型计算几何与材料参数（示意图，非比例绘制）", fontsize=11, y=1.02)
    fig.tight_layout()
    return save_figure(fig, output_dir / "figure_01_model_schematic")


def draw_response_figure(evidence: Evidence, variant: str, output_dir: Path) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(7.4, 2.65), sharex=True)
    f = evidence.fenicsx[variant]
    s = evidence.simpeg[variant]
    for ax, component in zip(axes, COMPONENTS):
        t = f.times
        ax.plot(t, np.abs(component_values(f, "reference", component)), color="#222222", lw=2.2, label="empymod参考")
        ax.plot(t, np.abs(component_values(f, "solver", component)), color="#1f77b4", lw=1.5, marker="o", ms=2.8, label="FEniCSx")
        ax.plot(t, np.abs(component_values(s, "solver", component)), color="#d55e00", lw=1.5, marker="s", ms=2.5, label="SimPEG")
        ax.set_xscale("log")
        ax.set_yscale("log")
        ax.set_xlabel("关断后时间 (s)")
        ax.set_ylabel(COMPONENT_LABELS[component])
        ax.grid(which="both", alpha=0.18)
        ax.set_title(COMPONENT_CN[component])
    axes[0].legend(loc="best")
    title = "无极化响应" if variant == "noip" else "极化响应"
    fig.suptitle(f"{title}：三种方法逐时刻对比", fontsize=11)
    fig.tight_layout()
    return save_figure(fig, output_dir / f"figure_{'03' if variant == 'noip' else '05'}_{variant}_responses")


def draw_error_figure(evidence: Evidence, variant: str, output_dir: Path) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(7.4, 2.65), sharex=True, sharey=True)
    f = evidence.fenicsx[variant]
    s = evidence.simpeg[variant]
    for ax, component in zip(axes, COMPONENTS):
        reference_f = component_values(f, "reference", component)
        reference_s = component_values(s, "reference", component)
        floor_f = evidence.summary["comparisons"][f"fenicsx_{variant}"][component]["floor"]
        floor_s = evidence.summary["comparisons"][f"simpeg_{variant}"][component]["floor"]
        err_f = 100 * relative_error(component_values(f, "solver", component), reference_f, floor_f)
        err_s = 100 * relative_error(component_values(s, "solver", component), reference_s, floor_s)
        ax.plot(f.times, err_f, color="#1f77b4", lw=1.8, marker="o", ms=2.8, label="FEniCSx")
        ax.plot(s.times, err_s, color="#d55e00", lw=1.8, marker="s", ms=2.5, label="SimPEG")
        ax.axhline(5.0, color="#555555", lw=1.2, ls=":", label="5%门槛")
        if variant == "ip" and component == "Ex":
            idx = int(np.argmax(err_s))
            ax.scatter(s.times[idx], err_s[idx], s=55, facecolor="none", edgecolor="#b00020", lw=1.4, zorder=5)
            ax.annotate("失败：7.421%", (s.times[idx], err_s[idx]), xytext=(1.75e-5, 6.8), fontsize=8, color="#b00020", arrowprops={"arrowstyle": "->", "color": "#b00020", "lw": 0.8})
        ax.set_xscale("log")
        ax.set_ylim(0, 8.0 if variant == "ip" else 5.4)
        ax.set_xlabel("关断后时间 (s)")
        ax.set_title(COMPONENT_CN[component])
        ax.grid(which="both", alpha=0.18)
    axes[0].set_ylabel("相对误差 (%)")
    axes[0].legend(loc="upper center", ncol=1)
    title = "无极化误差" if variant == "noip" else "极化误差"
    fig.suptitle(f"{title}：相对empymod有限线源参考", fontsize=11)
    fig.tight_layout()
    return save_figure(fig, output_dir / f"figure_{'04' if variant == 'noip' else '06'}_{variant}_errors")


def effect_percent(ip: np.ndarray, noip: np.ndarray) -> np.ndarray:
    if np.any(np.isclose(noip, 0.0, rtol=0.0, atol=1.0e-30)):
        raise AssertionError("Cannot form IP/no-IP effect with a zero denominator")
    return (ip / noip - 1.0) * 100.0


def draw_ip_effect(evidence: Evidence, output_dir: Path) -> Path:
    fig, axes = plt.subplots(1, 3, figsize=(7.4, 2.7), sharex=True)
    for ax, component in zip(axes, COMPONENTS):
        f_no = evidence.fenicsx["noip"]
        f_ip = evidence.fenicsx["ip"]
        s_no = evidence.simpeg["noip"]
        s_ip = evidence.simpeg["ip"]
        exact = effect_percent(
            component_values(f_ip, "reference", component),
            component_values(f_no, "reference", component),
        )
        f_eff = effect_percent(
            component_values(f_ip, "solver", component),
            component_values(f_no, "solver", component),
        )
        s_eff = effect_percent(
            component_values(s_ip, "solver", component),
            component_values(s_no, "solver", component),
        )
        ax.plot(f_no.times, exact, color="#222222", lw=2.2, label="empymod参考")
        ax.plot(f_no.times, f_eff, color="#1f77b4", lw=1.7, marker="o", ms=2.8, label="FEniCSx")
        ax.plot(s_no.times, s_eff, color="#d55e00", lw=1.7, marker="s", ms=2.5, label="SimPEG")
        ax.axhline(0, color="#777777", lw=0.8)
        ax.set_xscale("log")
        ax.set_xlabel("关断后时间 (s)")
        ax.set_ylabel("极化效应 (%)")
        ax.set_title(COMPONENT_CN[component])
        ax.grid(which="both", alpha=0.18)
    axes[0].legend(loc="best")
    fig.suptitle("极化效应：(IP/noIP - 1) × 100%", fontsize=11)
    fig.tight_layout()
    return save_figure(fig, output_dir / "figure_07_ip_effect")


def draw_max_error_summary(evidence: Evidence, output_dir: Path) -> Path:
    labels = ["FEniCSx\n无极化", "FEniCSx\n极化", "SimPEG\n无极化", "SimPEG\n极化"]
    keys = ["fenicsx_noip", "fenicsx_ip", "simpeg_noip", "simpeg_ip"]
    x = np.arange(len(labels), dtype=float)
    width = 0.22
    colors = {"Ex": "#4c78a8", "Hz": "#72b7b2", "dBzdt": "#f58518"}
    fig, ax = plt.subplots(figsize=(7.2, 3.3))
    for index, component in enumerate(COMPONENTS):
        values = [100 * evidence.summary["comparisons"][key][component]["max_relative_error"] for key in keys]
        bars = ax.bar(x + (index - 1) * width, values, width, label=COMPONENT_CN[component], color=colors[component])
        for bar, value in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width() / 2, value + 0.12, f"{value:.2f}", ha="center", va="bottom", fontsize=7.5)
    ax.axhline(5.0, color="#333333", lw=1.3, ls=":", label="5%门槛")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("最大相对误差 (%)")
    ax.set_ylim(0, 8.3)
    ax.set_title("三维离散方法相对empymod参考的最大误差")
    ax.grid(axis="y", alpha=0.2)
    ax.legend(ncol=4, loc="upper left")
    fig.tight_layout()
    return save_figure(fig, output_dir / "figure_08_max_error_summary")


def generate_figures(evidence: Evidence) -> list[Path]:
    configure_matplotlib()
    figures_dir = WORK_DIR / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    overview_copy = figures_dir / "figure_02_cross_code_overview.png"
    shutil.copy2(OVERVIEW_PATH, overview_copy)
    figures = [
        draw_model_schematic(evidence, figures_dir),
        overview_copy,
        draw_response_figure(evidence, "noip", figures_dir),
        draw_error_figure(evidence, "noip", figures_dir),
        draw_response_figure(evidence, "ip", figures_dir),
        draw_error_figure(evidence, "ip", figures_dir),
        draw_ip_effect(evidence, figures_dir),
        draw_max_error_summary(evidence, figures_dir),
    ]
    for path in figures:
        if not path.exists() or path.stat().st_size == 0:
            raise AssertionError(f"Missing figure: {path}")
    (WORK_DIR / "figure_contract.json").write_text(
        json.dumps(FIGURE_CONTRACT, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return figures


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None):
    rows = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[index], "365F91")
        if widths:
            table.rows[0].cells[index].width = Cm(widths[index])
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    header_properties.append(OxmlElement("w:cantSplit"))
    for row_index, row in enumerate(rows):
        table_row = table.add_row()
        table_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = table_row.cells
        for column_index, value in enumerate(row):
            set_cell_text(cells[column_index], value)
            if widths:
                cells[column_index].width = Cm(widths[column_index])
            if row_index % 2 == 1:
                set_cell_shading(cells[column_index], "EEF3F8")
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)


def add_figure(document: Document, path: Path, caption: str, width_cm: float = 16.6) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(document, caption)


def add_field(run, field_code: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name, size, color in [
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "365F91"),
        ("Heading 3", 11, "4F81BD"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    caption = styles["Caption"]
    caption.font.name = "Microsoft YaHei"
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor(60, 60, 60)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    if "Key conclusion" not in styles:
        key_style = styles.add_style("Key conclusion", WD_STYLE_TYPE.PARAGRAPH)
        key_style.font.name = "Microsoft YaHei"
        key_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        key_style.font.size = Pt(11)
        key_style.font.bold = True
        key_style.font.color.rgb = RGBColor.from_string("17365D")
        key_style.paragraph_format.left_indent = Cm(0.8)
        key_style.paragraph_format.right_indent = Cm(0.8)
        key_style.paragraph_format.space_before = Pt(8)
        key_style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "Song文献模型三维时域正演跨代码验证"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(110, 110, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    add_field(footer.add_run(), "PAGE")
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 110, 110)


def add_title_page(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Song文献均匀半空间模型\nFEniCSx、empymod与SimPEG\n三维时域正演对比报告")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("无极化与极化配对验证 | 关断后10 μs-1 ms")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("365F91")
    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("基于已完成的本地计算产物与机器可读审计结果生成\n不包含尚未计算的真实堤坝装置方案")
    note.runs[0].font.size = Pt(10.5)
    note.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(50)
    date_p.add_run(datetime.now().strftime("%Y年%m月%d日"))
    document.add_page_break()


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run(text)


def format_percent(value: float) -> str:
    return f"{100 * value:.6f}%"


def format_time(seconds: float) -> str:
    if seconds < 1e-3:
        return f"{seconds * 1e6:.3f} μs"
    return f"{seconds * 1e3:.3f} ms"


def metric_rows(evidence: Evidence, variant_key: str) -> list[list[str]]:
    rows = []
    for component in COMPONENTS:
        metric = evidence.summary["comparisons"][variant_key][component]
        rows.append(
            [
                COMPONENT_CN[component],
                format_percent(metric["max_relative_error"]),
                format_time(metric["time_at_max_s"]),
                format_percent(metric["rms_relative_error"]),
                format_percent(metric["terminal_relative_error"]),
                "通过" if metric["passed"] else "失败",
            ]
        )
    return rows


def add_method_parameter_tables(document: Document, evidence: Evidence) -> None:
    document.add_heading("4 三种方法的参数设置", level=1)
    document.add_paragraph(
        "本节将共同物理参数与三种方法自身的数值参数分开列出。所有数值均来自已保存的配置、诊断或结果文件；没有证据支持的参数不在表中补写。"
    )

    document.add_heading("4.1 共同物理与观测参数", level=2)
    add_table(
        document,
        ["参数", "取值", "说明"],
        [
            ["坐标约定", "z向下", "源和接收点均位于地表下0.1 m"],
            ["有限线源", "(-500, 0, 0.1)至(500, 0, 0.1) m", "长度1000 m，沿x方向"],
            ["发射电流", "10 A", "理想step-off，关断时间为0"],
            ["接收点", "(0, -500, 0.1) m", "距线源中点500 m"],
            ["空气电阻率", "10⁶ Ω·m", "准静态空气近似"],
            ["0-300 m电阻率", "100 Ω·m", "无极化与极化模型相同"],
            ["300 m以下电阻率", "100 Ω·m", "无极化背景"],
            ["Cole-Cole参数", "m=0.3，τ=1 s，c=0.3", "仅0-300 m参数区启用"],
            ["正式分量", "Ex、Hz、dBz/dt", "Ey仅作横向对称性诊断"],
            ["观测时间", "10 μs-1 ms，21个对数点", "本报告唯一正式验收时窗"],
            ["误差门槛", "5%", "固定门槛，未因结果改变"],
            ["误差分母下限", "各分量参考峰值的1%", "跨代码汇总采用相同规则"],
        ],
        widths=[4.1, 6.2, 6.1],
    )

    document.add_heading("4.2 FEniCSx参数", level=2)
    f_mesh = evidence.fenicsx_mesh["noip"]["preflight"]
    add_table(
        document,
        ["类别", "实际取值", "作用或证据"],
        [
            ["方程形式", "E场H(curl)形式", "主未知量为电场"],
            ["有限元", "二阶Nedelec边元", "保持切向连续性"],
            ["网格", f"{f_mesh['global_cells']:,}个全局四面体", "非结构Gmsh/DOLFINx网格"],
            ["全局自由度", f"{f_mesh['global_nedelec_dofs']:,}", "二阶边元自由度"],
            ["计算域", "x、y、空气高度和地下深度均扩展至±25 km", "降低人工边界影响"],
            ["有效加密区", "半径1000 m，深度797.885 m，目标尺寸80 m", "满足1 ms扩散尺度审计"],
            ["接收点局部网格", "5 m；加密半径60 m", "降低点接收空间误差"],
            ["线源进入网格", "88个四面体区间，264个积分点", "exact_tetra_intervals；覆盖率100%"],
            ["源投影", "charge_conserving", "端点平衡残差约1.24×10⁻¹³"],
            ["时间推进", "后向Euler，θ=1", "21个观测点之间每段8个子步"],
            ["内部时间步", "176步", "首个10 μs前至少16步"],
            ["线性求解器", "PETSc GMRES + HYPRE AMS", "rtol=10⁻⁸，atol=10⁻¹²，max_it=1000"],
            ["边界条件", "远边界PEC", "n×E=0"],
            ["Hz恢复", "H0 + Faraday积分", "H0采用高阶/自适应Biot-Savart恢复"],
            ["dBz/dt恢复", "-curl(E)点接收", "对应本基准正式磁感应变化率"],
            ["磁初始化求积", "名义q=8，审计q=2/4/6/8/10", "含近奇异单元自适应处理"],
            ["极化离散", "8项Debye近似", "相对L2拟合误差0.452638%，门槛1%"],
            ["并行规模", "12个MPI进程", "CPU利用率约1198%-1199%，无swap"],
            ["运行时间", "无极化1:28:34；极化1:28:48", "每个算例176个瞬变步"],
        ],
        widths=[3.3, 6.8, 6.3],
    )

    document.add_heading("4.3 empymod参考参数", level=2)
    f_cfg = evidence.fenicsx_config["noip"]
    add_table(
        document,
        ["类别", "实际取值", "说明"],
        [
            ["角色", "有限线源分层半空间参考", "不是三维网格离散求解器"],
            ["软件版本", "empymod 2.5.4", "由FEniCSx验证报告记录"],
            ["控制方程", "quasistatic", "水平/垂直介电常数设为0"],
            ["磁导率", "水平=1，垂直=1", "相对磁导率"],
            ["有限线源积分", f"主参考{f_cfg['empymod_srcpts']}点", "沿1000 m线源积分"],
            ["线源审计", f"{f_cfg['reference_audit_srcpts']}点", "9点与17点参考自收敛门槛0.5%"],
            ["Hankel变换", "DLF，key_201_2009", "pts_per_dec=0"],
            ["Fourier变换", "DLF，key_201_2012", "pts_per_dec=0"],
            ["参考时间", "与21个观测点完全一致", "不做响应时间平移"],
            ["极化模型", "精确Cole-Cole分层响应", "用于检验Debye近似后的三维求解器"],
            ["参考自审计", "Ex、Hz、dBz/dt均通过0.5%门槛", "17点有限线源积分作为高阶审计"],
        ],
        widths=[3.5, 6.2, 6.7],
    )

    document.add_heading("4.4 SimPEG参数", level=2)
    s_no = evidence.simpeg_json["noip"]
    s_ip = evidence.simpeg_json["ip"]
    mesh = s_no["mesh_stats"]
    dc_solver = s_no["initialization_solver_diagnostics"][0]
    amp_solver = s_no["initialization_solver_diagnostics"][1]
    add_table(
        document,
        ["类别", "实际取值", "作用或证据"],
        [
            ["离散方法", "拟合有限体积", "与FEniCSx边元离散独立"],
            ["空间级别", "S1 / 边界级别B0", "独立网格收敛后的首个通过级别"],
            ["网格规模", f"{mesh['n_cells']:,}个单元，{mesh['n_edges']:,}条边", "轴向单元数64×53×47"],
            ["计算域", "x、y、z约±25 km", "与物理边界审计一致"],
            ["时间细分", "T16", "每个观测区间16个子步"],
            ["瞬变求解次数", "336次/算例", "21个观测点对应内部推进"],
            ["DC初始化", f"{dc_solver['ksp_type'].upper()} + {dc_solver['pc_type']}", f"{dc_solver['backend_iterations']}次迭代，真残差{dc_solver['external_true_relative_residual']:.3e}"],
            ["Ampere初始化", f"{amp_solver['ksp_type'].upper()} + {amp_solver['pc_type']}", f"{amp_solver['backend_iterations']}次迭代，真残差{amp_solver['external_true_relative_residual']:.3e}"],
            ["初始化容差", "外部10⁻⁸，内部10⁻¹¹，max_it=4000", "初始化残差门槛未放宽"],
            ["无极化瞬变", f"最大迭代{s_no['linear_solver_summary']['maximum_backend_iterations']}，最大真残差{s_no['linear_solver_summary']['maximum_external_true_relative_residual']:.3e}", "全部reason=2"],
            ["极化瞬变", f"最大迭代{s_ip['linear_solver_summary']['maximum_backend_iterations']}，最大真残差{s_ip['linear_solver_summary']['maximum_external_true_relative_residual']:.3e}", "全部reason=2"],
            ["极化离散", "16项Debye近似", "相对L2拟合误差0.010243%，幅值均为正"],
            ["计算时间", "无极化2:56:10；极化2:49:14", "结果文件记录的compute elapsed"],
        ],
        widths=[3.5, 6.4, 6.5],
    )


def build_document(evidence: Evidence, figures: list[Path]) -> Path:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    document.core_properties.title = "Song文献模型FEniCSx、empymod与SimPEG三维时域正演对比报告"
    document.core_properties.subject = "10 μs-1 ms无极化与极化跨代码验证"
    document.core_properties.author = "Codex - evidence-backed local report"
    document.core_properties.keywords = "FEniCSx, empymod, SimPEG, TDEM, IP, dB/dt"
    document.core_properties.comments = "Generated from immutable local validation artifacts."

    add_title_page(document)

    document.add_heading("1 结论摘要", level=1)
    document.add_paragraph(
        "本报告整理已经完成的Song文献均匀半空间配对算例，对FEniCSx、empymod有限线源参考和SimPEG三种方法在关断后10 μs-1 ms时窗内的响应进行统一对比。正式分量为Ex、Hz和dBz/dt；Ey在该对称几何中理论接近零，仅作为数值对称性诊断。"
    )
    document.add_paragraph(
        "FEniCSx无极化最大误差分别为1.036429%、1.412497%和3.301669%；极化最大误差分别为1.993756%、1.090690%和2.687928%，全部低于固定5%门槛。",
        style="Key conclusion",
    )
    document.add_paragraph(
        "SimPEG无极化三个分量均通过。极化结果中Hz和dBz/dt通过，但Ex在10 μs处达到7.420635%，超过5%门槛；该失败点被完整保留，没有通过修改门槛、平滑或删除时间点进行掩盖。"
    )
    add_bullet(document, "对感应线圈而言，本基准中最直接相关的正式量是dBz/dt，单位T/s。")
    add_bullet(document, "Hz用于检查初始磁场、Faraday一致性和独立磁场恢复，不代表感应线圈在通电期间测量了静态磁场。")
    add_bullet(document, "结论仅覆盖文献模型与1 ms时窗，不包含真实堤坝、实际供电电缆、有限面积线圈或现场噪声。")

    document.add_heading("2 验证对象与验收规则", level=1)
    document.add_paragraph(
        "验证对象是1 km有限接地线源在100 Ω·m背景中的理想step-off响应。极化模型仅在0-300 m深度区间设置Cole-Cole参数m=0.3、τ=1 s、c=0.3，但该区间的直流电阻率仍为100 Ω·m，因此无极化和极化差异能够归因于频散/记忆效应，而不是简单更换背景电阻率。"
    )
    add_table(
        document,
        ["规则", "本报告采用值", "解释"],
        [
            ["正式时窗", "10 μs-1 ms", "21个对数时间点，含两端点"],
            ["正式分量", "Ex、Hz、dBz/dt", "三者均需独立通过"],
            ["诊断分量", "Ey", "对称参考近零，不使用普通相对误差验收"],
            ["参考", "empymod有限线源分层解", "主参考9点，17点进行线源求积审计"],
            ["误差门槛", "最大相对误差≤5%", "预先固定，没有事后放宽"],
            ["完整性", "不平移、不缩放、不平滑、不删点", "所有21点保留"],
        ],
        widths=[4.0, 5.5, 7.0],
    )

    document.add_heading("3 正演模型与观测量", level=1)
    add_figure(
        document,
        figures[0],
        "图1  Song文献模型的平面几何和垂向材料参数。该图只表示本次实际计算模型，非比例绘制。",
    )
    document.add_paragraph(
        "关断前，接地线源在地下建立稳定电流。理想step-off后，外加源电流归零，而地下场作为初始状态继续扩散。FEniCSx和SimPEG都推进离散场状态；empymod则给出相同有限线源和分层本构条件下的参考响应。"
    )
    document.add_paragraph("三项正式观测量定义为：")
    add_bullet(document, "Ex：接收位置沿x方向的电场，单位V/m。")
    add_bullet(document, "Hz：垂直磁场强度，单位A/m；用于磁初始化和Faraday积分路径验证。")
    add_bullet(document, "dBz/dt：垂直磁通密度时间导数，单位T/s；满足dB/dt=-curl(E)。")
    document.add_paragraph(
        "对于法向为z的小型感应线圈，输出电压近似满足V=-N Aeff dBz/dt。因此，本报告中的dBz/dt是与关断后感应电压最直接对应的算法量。报告中的Hz是数值状态/验证量，不意味着本次装置在供电期间测量静态磁场。"
    )

    document.add_heading("3.1 Song文献原图及与本次算例的对应关系", level=2)
    document.add_paragraph(
        "以下四幅图直接摘自Song等（2025）原文，用于说明论文中的SOTEM装置、三维模型及极化响应。"
        "它们与前面的报告重绘图用途不同：文献原图用于追溯模型出处，报告图1用于清楚表达本次实际计算坐标和参数。"
        "文献原图均保留论文图号，并在本报告题注中给出出处；不得将其视为FEniCSx、empymod或SimPEG的新计算结果。"
    )
    add_figure(
        document,
        SONG_PAPER_FIGURES["layout"],
        "文献原图1  Song等（2025）图1：SOTEM方法示意图。"
        "该图是论文的通用装置示意，不是本项目真实堤坝装置。"
        "来源：Journal of Applied Geophysics 233 (2025) 105613，"
        "DOI: 10.1016/j.jappgeo.2024.105613。",
        width_cm=11.5,
    )
    document.add_page_break()
    add_figure(
        document,
        SONG_PAPER_FIGURES["model"],
        "文献原图2  Song等（2025）图4：三维正演模型网格，中间层为极化层。"
        "本次交叉验证采用该组层状极化模型框架，并通过算例文件固定实际坐标和参数。"
        "来源同上。",
        width_cm=13.5,
    )
    document.add_page_break()
    add_figure(
        document,
        SONG_PAPER_FIGURES["ex"],
        "文献原图3  Song等（2025）图7：有极化与无极化条件下的Ex响应及相对极化效应。"
        "论文图覆盖10 μs-1 s；本报告的正式数值验收只覆盖10 μs-1 ms。"
        "来源同上。",
        width_cm=16.2,
    )
    document.add_page_break()
    add_figure(
        document,
        SONG_PAPER_FIGURES["hz"],
        "文献原图4  Song等（2025）图8：有极化与无极化条件下的Hz响应及相对极化效应。"
        "论文图覆盖10 μs-1 s；本报告的正式数值验收只覆盖10 μs-1 ms。"
        "来源同上。",
        width_cm=16.2,
    )
    document.add_paragraph(
        "文献图7和图8用于核对响应形态、符号变化和极化效应随时间的总体规律，不作为本报告逐点误差的数字化数据源。"
        "本报告的empymod参考值由同一有限线源、层状电阻率和Cole-Cole参数独立计算。"
        "此外，Song文献图7和图8直接给出Ex与Hz；与感应线圈电压成正比的dBz/dt由Faraday关系独立计算并在后文对比。"
    )

    add_method_parameter_tables(document, evidence)

    document.add_heading("5 跨代码总览", level=1)
    add_figure(
        document,
        figures[1],
        "图2  10 μs-1 ms三种方法总览。左列为绝对响应，右列为相对empymod参考的误差；虚线表示极化，点线表示5%门槛。",
        width_cm=16.8,
    )
    document.add_paragraph(
        "总览图显示三种方法在响应形态上整体一致。无极化和极化曲线均随时间演化，但并非单一指数函数；无限/半无限扩散问题是连续空间模态的叠加，晚期通常表现为幂律型衰减。误差图则将形态相似性转化为固定门槛下的定量判断。"
    )

    document.add_heading("6 无极化结果", level=1)
    add_figure(document, figures[2], "图3  无极化Ex、Hz和dBz/dt响应。所有21个观测点均保留。")
    add_figure(document, figures[3], "图4  无极化逐时刻相对误差。FEniCSx和SimPEG三个正式分量均低于5%。")
    add_table(
        document,
        ["方法/分量", "最大误差", "最大误差时间", "RMS误差", "1 ms误差", "结果"],
        [["FEniCSx " + row[0], *row[1:]] for row in metric_rows(evidence, "fenicsx_noip")]
        + [["SimPEG " + row[0], *row[1:]] for row in metric_rows(evidence, "simpeg_noip")],
        widths=[3.4, 2.5, 3.0, 2.5, 2.5, 2.0],
    )
    document.add_paragraph(
        "FEniCSx无极化最大误差出现在1 ms：Ex为1.036429%、Hz为1.412497%、dBz/dt为3.301669%。SimPEG无极化最大误差分别为3.210106%、0.630309%和3.452765%，均通过。FEniCSx在Ex上更接近参考，SimPEG在Hz上更接近参考，dBz/dt精度相近。"
    )

    document.add_heading("7 极化结果", level=1)
    add_figure(document, figures[4], "图5  极化Ex、Hz和dBz/dt响应。极化改变了三个分量的幅值和时间演化。")
    add_figure(document, figures[5], "图6  极化逐时刻相对误差。SimPEG Ex在10 μs处为7.420635%，明确标记失败。")
    add_table(
        document,
        ["方法/分量", "最大误差", "最大误差时间", "RMS误差", "1 ms误差", "结果"],
        [["FEniCSx " + row[0], *row[1:]] for row in metric_rows(evidence, "fenicsx_ip")]
        + [["SimPEG " + row[0], *row[1:]] for row in metric_rows(evidence, "simpeg_ip")],
        widths=[3.4, 2.5, 3.0, 2.5, 2.5, 2.0],
    )
    document.add_paragraph(
        "FEniCSx极化三个正式分量均通过：Ex最大误差1.993756%出现在19.953 μs，Hz和dBz/dt最大误差分别为1.090690%和2.687928%，均出现在1 ms。SimPEG极化Hz和dBz/dt通过，但Ex在10 μs处为7.420635%。"
    )
    document.add_paragraph(
        "该SimPEG失败不能归因于线性求解器未收敛：全部336个瞬变求解均为正收敛reason=2，最大真残差3.573×10⁻¹¹；16项Debye拟合相对L2误差仅0.010243%。既有T8/T16/T32审计表明Ex随时间细分改善有限，剩余偏差主要属于S1空间/共模误差。"
    )

    document.add_heading("8 极化效应本身的恢复精度", level=1)
    add_figure(document, figures[6], "图7  三种正式分量的极化效应曲线，定义为(IP/noIP - 1)×100%。")
    effect = evidence.summary["ip_effect_percentage_point_error"]
    add_table(
        document,
        ["方法", "Ex效应最大误差", "Hz效应最大误差", "dBz/dt效应最大误差"],
        [
            ["FEniCSx", f"{effect['fenicsx']['Ex']['max_abs_percentage_point_error']:.6f}个百分点", f"{effect['fenicsx']['Hz']['max_abs_percentage_point_error']:.6f}个百分点", f"{effect['fenicsx']['dBzdt']['max_abs_percentage_point_error']:.6f}个百分点"],
            ["SimPEG", f"{effect['simpeg']['Ex']['max_abs_percentage_point_error']:.6f}个百分点", f"{effect['simpeg']['Hz']['max_abs_percentage_point_error']:.6f}个百分点", f"{effect['simpeg']['dBzdt']['max_abs_percentage_point_error']:.6f}个百分点"],
        ],
        widths=[3.2, 4.4, 4.4, 4.4],
    )
    document.add_paragraph(
        "FEniCSx对Ex极化效应的最大百分点误差为1.356574，低于SimPEG的2.525546；SimPEG对Hz极化效应略优；dBz/dt两者接近。这说明只比较IP绝对响应还不够，还必须比较IP相对于noIP的差异是否被正确恢复。"
    )

    document.add_heading("9 最大误差与方法间差异", level=1)
    add_figure(document, figures[7], "图8  FEniCSx和SimPEG在无极化/极化情况下的最大误差汇总。5%线为固定验收门槛。")
    cross = evidence.summary["comparisons"]
    add_table(
        document,
        ["模型", "Ex最大方法间差异", "Hz最大方法间差异", "dBz/dt最大方法间差异"],
        [
            ["无极化", format_percent(cross["fenicsx_vs_simpeg_noip"]["Ex"]["max_exact_scaled_difference"]), format_percent(cross["fenicsx_vs_simpeg_noip"]["Hz"]["max_exact_scaled_difference"]), format_percent(cross["fenicsx_vs_simpeg_noip"]["dBzdt"]["max_exact_scaled_difference"])],
            ["极化", format_percent(cross["fenicsx_vs_simpeg_ip"]["Ex"]["max_exact_scaled_difference"]), format_percent(cross["fenicsx_vs_simpeg_ip"]["Hz"]["max_exact_scaled_difference"]), format_percent(cross["fenicsx_vs_simpeg_ip"]["dBzdt"]["max_exact_scaled_difference"])],
        ],
        widths=[3.0, 4.5, 4.5, 4.5],
    )
    document.add_paragraph(
        "极化Ex的FEniCSx-SimPEG最大方法间差异为5.764285%，出现在10 μs。该差异超过5%的原因是SimPEG极化Ex自身在该点未通过；FEniCSx同一点相对参考的误差仍低于2%。因此不能把该差异解释为两个三维算法都失败。"
    )

    document.add_heading("10 数值可信性证据", level=1)
    add_table(
        document,
        ["检查项", "证据", "判断"],
        [
            ["有限线源积分", "88段、264积分点、100%覆盖、无间隙无重叠", "通过"],
            ["参考线源自收敛", "empymod主参考9点、审计17点，三分量均低于0.5%", "通过"],
            ["FEniCSx空间分辨率", "221,164四面体、1,413,482边元自由度；内部300 m界面已加密", "通过本时窗"],
            ["FEniCSx时间离散", "后向Euler，21个观测点、176个内部步；T8审计", "通过本时窗"],
            ["扩散边界", "1 ms推荐解析半径/深度797.885 m；实际解析半径1000 m", "通过"],
            ["磁初始化", "自适应Biot-Savart与q=2/4/6/8/10审计；近奇异单元专门处理", "通过"],
            ["MPI一致性", "多进程几何索引、集体装配和Biot贡献归约已修复；12进程无swap", "通过"],
            ["SimPEG初始化", "DC和Ampere初始化真残差均低于10⁻⁸", "通过"],
            ["SimPEG瞬变", "无极化/极化各336次，全部reason=2", "通过"],
            ["极化材料拟合", "FEniCSx 8项Debye误差0.452638%；SimPEG 16项误差0.010243%", "均通过各自门槛"],
        ],
        widths=[3.7, 9.5, 3.2],
    )

    document.add_heading("11 适用边界与不能外推的结论", level=1)
    document.add_paragraph(
        "本报告是均匀/分层文献模型的内部数值验证，不是现场工程验证。以下内容尚未由本次数据证明："
    )
    add_bullet(document, "不证明真实坝体、复杂地形、三维渗流通道或金属设施模型已经正确。")
    add_bullet(document, "不证明100 m上下游电极布置、坝顶/坝坡接收或真实供电电缆磁场已经建模。")
    add_bullet(document, "不证明有限面积线圈、线圈传递函数、关断斜坡、仪器死区和现场噪声已经纳入。")
    add_bullet(document, "不证明1 ms以后至1 s的长时窗已经通过网格、时间步和边界收敛。")
    add_bullet(document, "不将Ey永久设为零；只有当前对称基准中Ey作为数值泄漏诊断。")
    document.add_paragraph(
        "因此，当前可支持的结论是：我们的FEniCSx核心求解器能够在Song文献模型的10 μs-1 ms时窗内，正确恢复无极化和极化Ex、Hz及dBz/dt，并与独立有限线源参考和独立三维有限体积结果形成可审计的一致性证据。"
    )

    document.add_heading("12 复现与证据索引", level=1)
    add_table(
        document,
        ["证据", "标识"],
        [
            ["代码分支", "codex/lei2023-sotem-benchmark"],
            ["核心计算提交", "5d03d6ebc4beb2f26c5a92c4c535a9087d1b1589"],
            ["FEniCSx无极化数据SHA-256", evidence.summary["inputs"]["fenicsx_noip"]["sha256"]],
            ["FEniCSx极化数据SHA-256", evidence.summary["inputs"]["fenicsx_ip"]["sha256"]],
            ["SimPEG无极化数据SHA-256", evidence.summary["inputs"]["simpeg_noip"]["sha256"]],
            ["SimPEG极化数据SHA-256", evidence.summary["inputs"]["simpeg_ip"]["sha256"]],
            ["共同时间轴哈希", evidence.simpeg_json["noip"]["time_hash"]],
            ["SimPEG网格哈希", evidence.simpeg_json["noip"]["mesh_hash"]],
        ],
        widths=[5.0, 11.4],
    )
    document.add_paragraph(
        "机器可读汇总文件记录了每个分量的最大、RMS和末端误差，以及最大误差时刻和极化效应百分点误差。最终报告生成前会重新校验四个输入NPZ的SHA-256，任何结果文件变化都会使构建失败。"
    )

    document.add_heading("附录A 误差定义", level=1)
    document.add_paragraph(
        "逐点相对误差采用 |dnum-dref| / max(|dref|, dfloor)，其中dfloor为该分量参考峰值的1%。最大误差为21个正式时间点中的最大值，RMS误差为全部正式时间点相对误差的均方根。该分母下限仅防止参考过零附近的无意义放大，不改变原始响应。"
    )
    document.add_paragraph(
        "Ey的参考值在当前对称几何中接近数值零，因此不使用普通相对误差作为正式验收；报告只记录其相对主分量的泄漏量。"
    )

    temp_output = OUTPUT_PATH.with_suffix(".tmp.docx")
    document.save(temp_output)
    os.replace(temp_output, OUTPUT_PATH)
    return OUTPUT_PATH


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check-only", action="store_true", help="Validate evidence without writing figures or DOCX")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    evidence = load_evidence()
    if args.check_only:
        retained = [
            f"{variant}/{component}"
            for variant, metrics in evidence.summary["comparisons"].items()
            if variant in EXPECTED_MAX
            for component, item in metrics.items()
            if not item["passed"]
        ]
        print(
            "EVIDENCE CHECK PASS: 21 observations, four formal solver variants, "
            f"retained failures={retained}"
        )
        return 0
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    figures = generate_figures(evidence)
    output = build_document(evidence, figures)
    print(f"REPORT BUILT: {output}")
    print(f"FIGURES: {len(figures)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
