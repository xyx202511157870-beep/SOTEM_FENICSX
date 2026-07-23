#!/usr/bin/env python3
"""Verify structure and evidence-critical text in the Song comparison DOCX."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document


REQUIRED_TEXT = (
    "1.036429%",
    "1.412497%",
    "3.301669%",
    "1.993756%",
    "1.090690%",
    "2.687928%",
    "3.210106%",
    "0.630309%",
    "3.452765%",
    "7.420635%",
    "0.654283%",
    "3.445477%",
    "10 μs",
    "1 ms",
    "5%",
    "Ey",
    "仅作为数值对称性诊断",
    "不包含尚未计算的真实堤坝装置方案",
)

FORBIDDEN_DRAFT_MARKERS = ("<<replace", "[placeholder]", "待补充内容", "待插图")


def collect_text(document: Document) -> str:
    chunks: list[str] = []
    chunks.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def verify(path: Path) -> dict[str, int]:
    if not path.is_file() or path.stat().st_size < 100_000:
        raise AssertionError(f"DOCX missing or unexpectedly small: {path}")

    document = Document(path)
    text = collect_text(document)
    missing = [value for value in REQUIRED_TEXT if value not in text]
    if missing:
        raise AssertionError(f"Missing required report text: {missing}")
    present_markers = [value for value in FORBIDDEN_DRAFT_MARKERS if value in text]
    if present_markers:
        raise AssertionError(f"Draft markers remain: {present_markers}")

    figure_count = len(document.inline_shapes)
    table_count = len(document.tables)
    heading_count = sum(
        1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")
    )
    caption_count = sum(
        1 for paragraph in document.paragraphs if paragraph.style.name == "Caption"
    )

    if figure_count < 8:
        raise AssertionError(f"Expected at least 8 figures, found {figure_count}")
    if caption_count < 8:
        raise AssertionError(f"Expected at least 8 captions, found {caption_count}")
    if table_count < 9:
        raise AssertionError(f"Expected at least 9 tables, found {table_count}")
    if heading_count < 15:
        raise AssertionError(f"Expected at least 15 headings, found {heading_count}")

    props = document.core_properties
    if not props.title or not props.subject or not props.author:
        raise AssertionError("Core document properties are incomplete")
    if "真实堤坝" not in text or "1 s" not in text or "有限面积线圈" not in text:
        raise AssertionError("Engineering-scope limitations are incomplete")
    if "SimPEG" not in text or "empymod" not in text or "FEniCSx" not in text:
        raise AssertionError("One or more method names are absent")

    return {
        "figures": figure_count,
        "captions": caption_count,
        "tables": table_count,
        "headings": heading_count,
        "characters": len(text),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    stats = verify(args.docx)
    print("PASS " + " ".join(f"{key}={value}" for key, value in stats.items()))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        raise
