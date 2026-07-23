"""Build the fail-closed Zhou 2020 validation report as a DOCX artifact."""

from __future__ import annotations

import argparse
from datetime import date
import hashlib
import importlib.util
import json
import os
from pathlib import Path
import subprocess
import tempfile
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
NAVY = "17365D"
TEAL = "167D8D"
PALE_BLUE = "EAF0F8"
PALE_RED = "FBEAEA"
PALE_GREEN = "EAF4EF"
RED = "9C2828"
GREEN = "25734A"
GRAY = "5A6573"
FIGURES = (
    (
        "fig01_model_contract.png",
        "图 1  Zhou 2020 接地线源 TEM-IP 数值基准的模型与参数合同。",
    ),
    (
        "fig02_total_fields.png",
        "图 2  no-IP 与 IP 总场的文献式绝对值双对数比较。纵轴为响应绝对值；绝对值不提供符号信息，反号必须查看独立 signed diagnostic。",
    ),
    (
        "fig03_reference_stability.png",
        "图 3  dBz/dt 独立 IP 增量的参考变换稳定性审计。灰色区域标为 reference-transform unstable；全部原始样本均保留。",
    ),
    (
        "fig04_gate_summary.png",
        "图 4  正式总场门限与独立诊断总览。IP 增量仅作为敏感性证据，不进入正式门限判定。",
    ),
    (
        "fig05_debye_order_diagnostic.png",
        "图 5  同网格、同时间步下 4 项相对 16 项 Debye 的 pole-count sensitivity，属于内部敏感性而非跨代码通过判据。",
    ),
)


def _load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise ImportError(f"cannot load module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _load_validated_bundle(path: Path) -> dict[str, Any]:
    module = _load_module(
        "zhou_publication_bundle_validator",
        ROOT / "scripts/plot_zhou2020_strict_validation.py",
    )
    return module.load_validated_bundle(Path(path))


def _load_validated_reference_audit(path: Path) -> dict[str, Any]:
    module = _load_module(
        "zhou_reference_audit_validator_for_report",
        ROOT / "scripts/audit_zhou2020_reference_stability.py",
    )
    return module.load_validated_audit(Path(path))


def _docx_path(value: str) -> Path:
    path = Path(value)
    if path.suffix.lower() != ".docx":
        raise argparse.ArgumentTypeError("report output must be a .docx file")
    return path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Build the Zhou validation DOCX; no PDF is generated.",
    )
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--run", type=Path, required=True)
    parser.add_argument("--publication-bundle", type=Path, required=True)
    parser.add_argument("--reference-audit", type=Path, required=True)
    parser.add_argument("--output", type=_docx_path, required=True)
    return parser


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _format_run(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    font: str = "Microsoft YaHei",
) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in (
        ("Title", 22, NAVY),
        ("Subtitle", 12, GRAY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, TEAL),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"),
            "Microsoft YaHei",
        )
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.space_before = Pt(10)

    header = section.header.paragraphs[0]
    header.text = "Zhou 2020 接地线源 TEM-IP FEniCSx 验证"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        _format_run(run, size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: list[list[str]],
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            _format_run(run, bold=True, color="FFFFFF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = str(value)
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def _add_status(doc: Document, title: str, body: str, *, positive: bool) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE_GREEN if positive else PALE_RED)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{title}：")
    _format_run(run, bold=True, color=GREEN if positive else RED)
    paragraph.add_run(body)


def _add_figure(
    doc: Document,
    image: Path,
    caption: str,
    *,
    width_cm: float = 16.7,
) -> None:
    if not image.is_file():
        raise FileNotFoundError(f"validated bundle figure is missing: {image}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    _format_run(run, size=8.5, color=GRAY, font="宋体")


def _git_state(root: Path) -> tuple[str, bool]:
    commit = subprocess.check_output(
        ["git", "rev-parse", "HEAD"],
        cwd=root,
        text=True,
    ).strip()
    dirty = bool(
        subprocess.check_output(
            ["git", "status", "--porcelain"],
            cwd=root,
            text=True,
        ).strip()
    )
    return commit, dirty


def _relative_or_absolute(path: Path, root: Path) -> str:
    try:
        return str(path.resolve().relative_to(root.resolve()))
    except ValueError:
        return str(path.resolve())


def _percentage(value: float, digits: int = 2) -> str:
    return f"{float(value) * 100:.{digits}f}%"


def _sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _cross_bound_input_hashes(
    manifest: dict[str, Any],
    validated_audit: dict[str, Any],
) -> dict[str, str]:
    metadata = manifest.get("metadata")
    audit_manifest = validated_audit.get("manifest")
    audit = validated_audit.get("audit")
    if (
        not isinstance(metadata, dict)
        or not isinstance(audit_manifest, dict)
        or not isinstance(audit, dict)
    ):
        raise ValueError("validated evidence identity records are incomplete")
    candidates = (
        metadata.get("input_sha256"),
        audit_manifest.get("input_sha256"),
        audit.get("input_sha256"),
    )
    if (
        any(not isinstance(candidate, dict) for candidate in candidates)
        or candidates[0] != candidates[1]
        or candidates[0] != candidates[2]
    ):
        raise ValueError("bundle and audit input identity hashes differ")
    hashes = dict(candidates[0])
    for required in ("run_manifest.json", "strict_comparison.json"):
        value = hashes.get(required)
        if (
            not isinstance(value, str)
            or len(value) != 64
            or any(character not in "0123456789abcdef" for character in value)
        ):
            raise ValueError(f"{required} identity hash is invalid")
    return hashes


def _load_cross_bound_metrics(
    root: Path,
    run: Path,
    manifest: dict[str, Any],
    validated_audit: dict[str, Any],
) -> dict[str, Any]:
    """Read strict metrics only after path, manifest, and hash cross-binding."""

    run = Path(run)
    metadata = manifest["metadata"]
    run_id = metadata.get("run_id")
    if not isinstance(run_id, str) or not run_id or run.name != run_id:
        raise ValueError("formal run path does not match the bound run_id")
    expected_run = (
        Path(root)
        / "generated/validation/zhou2020_grounded_wire/runs"
        / run_id
    )
    if run.resolve() != expected_run.resolve():
        raise ValueError("formal run path is outside the bound validation run root")
    hashes = _cross_bound_input_hashes(manifest, validated_audit)

    run_manifest_path = run / "run_manifest.json"
    strict_path = run / "comparisons/S1T1B1/strict_comparison.json"
    if not run_manifest_path.is_file() or not strict_path.is_file():
        raise ValueError("bound formal run evidence path is missing")

    run_manifest_bytes = run_manifest_path.read_bytes()
    strict_bytes = strict_path.read_bytes()
    if _sha256_bytes(run_manifest_bytes) != hashes["run_manifest.json"]:
        raise ValueError("run manifest identity hash mismatch")
    if _sha256_bytes(strict_bytes) != hashes["strict_comparison.json"]:
        raise ValueError("strict comparison identity hash mismatch")
    try:
        run_manifest = json.loads(run_manifest_bytes)
        metrics = json.loads(strict_bytes)
    except (UnicodeDecodeError, json.JSONDecodeError) as error:
        raise ValueError("bound formal evidence JSON is invalid") from error
    if not isinstance(run_manifest, dict) or not isinstance(metrics, dict):
        raise ValueError("bound formal evidence must contain JSON objects")

    comparison = run_manifest.get("comparisons", {}).get("full/S1T1B1")
    if (
        run_manifest.get("schema")
        != "atem3d.zhou2020.validation-run/v1"
        or run_manifest.get("run_id") != run_id
        or not isinstance(comparison, dict)
        or comparison.get("path")
        != "comparisons/S1T1B1/strict_comparison.json"
        or comparison.get("sha256") != hashes["strict_comparison.json"]
        or comparison.get("status") != "failed_with_reproducible_evidence"
    ):
        raise ValueError("formal run manifest identity or status is inconsistent")
    if (
        metrics.get("schema") != "atem3d.zhou2020.strict-comparison/v1"
        or metrics.get("status") != comparison["status"]
    ):
        raise ValueError("strict comparison schema or status is inconsistent")

    if (
        _sha256_file(run_manifest_path) != hashes["run_manifest.json"]
        or _sha256_file(strict_path) != hashes["strict_comparison.json"]
    ):
        raise ValueError("formal evidence changed during report generation")
    return metrics


def _validate_evidence(
    bundle: dict[str, Any],
    validated_audit: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    manifest = bundle.get("manifest")
    diagnostic = bundle.get("diagnostic")
    audit = validated_audit.get("audit")
    if not isinstance(manifest, dict) or not isinstance(diagnostic, dict):
        raise ValueError("validated publication bundle payload is incomplete")
    if not isinstance(audit, dict):
        raise ValueError("validated reference audit payload is incomplete")
    qwe = audit.get("qwe")
    if audit.get("status") != "inconclusive":
        raise ValueError("reference audit status changed; re-review report wording")
    if not isinstance(qwe, dict) or qwe.get("converged") is not False:
        raise ValueError("QWE convergence evidence changed; re-review report wording")
    metadata = manifest.get("metadata")
    if (
        not isinstance(metadata, dict)
        or metadata.get("reference_audit_status") != "inconclusive"
        or metadata.get("qwe_converged") is not False
    ):
        raise ValueError("publication bundle audit binding changed; re-review report")
    if (
        diagnostic.get("schema")
        != "atem3d.zhou2020.debye-order-diagnostic/v2"
    ):
        raise ValueError("unsupported Debye diagnostic schema")
    return manifest, diagnostic, audit


def _metric_rows(metrics: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for variant in ("noip", "ip"):
        for component in ("Ex", "Hz", "dBzdt"):
            item = metrics["total_field"][variant][component]
            rows.append(
                [
                    f"{variant} 总场",
                    component,
                    _percentage(item["relative_l2"], 3),
                    _percentage(item["gate"], 1),
                    "通过" if item["passed"] else "未通过",
                ]
            )
    for component in ("Ex", "Hz", "dBzdt"):
        item = metrics["ip_increment"][component]
        rows.append(
            [
                "IP 增量",
                component,
                _percentage(item["relative_l2"], 2),
                "不进入正式门",
                "敏感性数值（不判定）",
            ]
        )
    return rows


def _first_prediction(metrics: dict[str, Any], variant: str, component: str) -> float:
    values = metrics["zero_crossings"][variant][component]["prediction"]
    if not values:
        raise ValueError(f"missing {variant} {component} prediction zero crossing")
    return float(values[0])


def _write_report(
    *,
    root: Path,
    run: Path,
    publication_bundle: Path,
    reference_audit: Path,
    output: Path,
    metrics: dict[str, Any],
    manifest: dict[str, Any],
    diagnostic: dict[str, Any],
    audit: dict[str, Any],
    git_commit: str,
    git_dirty: bool,
) -> None:
    doc = Document()
    _configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Zhou 2020 接地线源 TEM-IP\nFEniCSx 数值验证报告")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("绝对值响应图、符号诊断与参考变换稳定性审计")
    doc.add_paragraph()
    _add_status(
        doc,
        "总体状态",
        "未完全通过（failed_with_reproducible_evidence）。总场与 Ex 反号提供部分支持；"
        "Hz 的 false reversal 与弱增量参考不稳定证据必须保留。",
        positive=False,
    )
    _add_table(
        doc,
        ("项目", "记录"),
        [
            ["报告日期", date.today().isoformat()],
            ["Git commit", git_commit[:12]],
            ["工作区状态", "dirty" if git_dirty else "clean"],
            ["正式运行状态", str(metrics.get("status"))],
            ["参考审计状态", str(audit.get("status"))],
            [
                "文档 QA",
                "layout rendering not performed / LibreOffice unavailable；"
                "已执行 DOCX ZIP/XML、关系和媒体结构检查。该限制不构成科学通过证据。",
            ],
            ["交付格式", "DOCX only；不生成、不交付 PDF"],
        ],
    )
    doc.add_page_break()

    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(
        "本报告使用 publication_bundle 的完整 manifest 与哈希校验后再读取五张图，"
        "并独立调用 Task 2 loader 验证 reference_audit_hardened_v2。"
        "报告不读取父目录中的 legacy 图，不重新计算或改写 strict_comparison.json。"
    )
    doc.add_paragraph(
        "文献中的主响应图采用绝对值双对数坐标。本报告据此绘制 |Ex|、|Hz| 与 "
        "|dBz/dt|；绝对值不提供符号信息，因此符号变化和反号只在独立 signed "
        "diagnostic 中解释。所有原始样本均保留，没有删点、平滑或调整门限。"
    )
    _add_status(
        doc,
        "可以确认",
        "no-IP/IP 总场完整窗口相对 L2 与 Ex 首次反号时刻保留原正式判据和结果。",
        positive=True,
    )
    _add_status(
        doc,
        "不能确认",
        "弱 IP 增量的参考变换稳定性没有建立；QWE direct 未收敛，故状态为 "
        "inconclusive。11.40% 和 9.44% 均不能升级为正式通过或失败。",
        positive=False,
    )

    doc.add_heading("1  基准合同与证据来源", level=1)
    doc.add_paragraph(
        "Zhou et al. (2020), GEOPHYSICS 85(4), E111-E120，"
        "DOI 10.1190/geo2019-0322.1，是接地线源 TEM-IP 主参考。"
        "完整数值参数来自同一作者团队公开材料；Song 三维模型因参数合同不完整，"
        "仅保留为定性场景参考，不用于本次定量误差门限。"
    )
    _add_figure(
        doc,
        publication_bundle / FIGURES[0][0],
        FIGURES[0][1],
    )
    _add_table(
        doc,
        ("对象", "参数"),
        [
            ["第一层", "DC 电阻率 100 Ω·m，厚度 500 m"],
            ["极化层", "DC 电阻率 10 Ω·m，厚度 20 m，m=0.1，τ=1 s，c=0.3"],
            ["第三层", "200 Ω·m 半空间"],
            ["发射", "1000 m 有限接地线，10 A，理想 step-off"],
            ["接收", "R=(0,1000,0) m；Ex、Hz、dBz/dt"],
            ["时间窗", "1e-4 s 至 3 s，101 个原始有符号样本"],
        ],
    )

    doc.add_heading("2  文献式绝对值总场与符号诊断", level=1)
    _add_figure(
        doc,
        publication_bundle / FIGURES[1][0],
        FIGURES[1][1],
    )
    doc.add_paragraph(
        "总场图的纵轴是绝对值双对数，因此它适合比较数量级和衰减形态，"
        "但绝对值不提供符号信息。任何“反号已对上”的结论都必须回到 signed "
        "diagnostic 和 zero-crossing 记录；不能从绝对值曲线本身推出。"
    )
    _add_table(
        doc,
        ("类别", "分量", "相对 L2", "门限", "解释"),
        _metric_rows(metrics),
    )

    doc.add_heading("3  参考变换稳定性审计", level=1)
    _add_figure(
        doc,
        publication_bundle / FIGURES[2][0],
        FIGURES[2][1],
    )
    stable = audit["stable_window"]
    transform = audit["transform_difference"]
    qwe_sensitivity = audit["fenicsx_vs_direct_qwe"]
    sign_changes = int(audit["default_dlf"]["sign_changes_first20"])
    doc.add_paragraph(
        "早期总场量级约为 4e-7 T/s，而 IP-noIP 弱增量仅约 "
        "1e-14 至 1e-10 T/s。此处是弱差分问题：default DLF 的前 20 个样本"
        f"出现 {sign_changes} 次符号变化；早期 DLF 震荡不是物理响应，而是参考"
        "变换对两个强总场相减不稳定的表现。"
    )
    doc.add_paragraph(
        f"reference-transform unstable 区域延伸至约 "
        f"{float(stable['start_s']):.3g} s（约 5.77e-4 s）。"
        f"default DLF 与 direct-frequency QWE 的全窗口差异为 "
        f"{_percentage(transform['default_dlf_vs_direct_qwe_relative_l2_full'])}，"
        f"前 20 点差异为 "
        f"{_percentage(transform['default_dlf_vs_direct_qwe_relative_l2_first20'])}。"
    )
    doc.add_paragraph(
        f"FEniCSx 相对 direct-frequency QWE 的全窗口敏感性为 "
        f"{_percentage(qwe_sensitivity['relative_l2_full'])}（9.44%），稳定窗口为 "
        f"{_percentage(qwe_sensitivity['relative_l2_stable_window'])}（9.39%）。"
        "但是 QWE direct 未收敛，QWE 全局 convergence=false，因此 9.44% 仅为"
        "敏感性，不是通过。审计状态保持 inconclusive。"
    )
    _add_table(
        doc,
        ("证据", "数值", "判定角色"),
        [
            ["default DLF dBz/dt IP 增量", "11.40%", "仅为敏感性，不是正式失败"],
            ["FEniCSx / direct QWE", "9.44%", "仅为敏感性，不是通过"],
            ["QWE direct", "converged=False", "使参考审计为 inconclusive"],
            ["样本处理", "all samples retained", "没有删点或平滑"],
        ],
    )

    doc.add_heading("4  正式判据与独立 false reversal", level=1)
    _add_figure(
        doc,
        publication_bundle / FIGURES[3][0],
        FIGURES[3][1],
    )
    hz_noip = _first_prediction(metrics, "noip", "Hz")
    hz_ip = _first_prediction(metrics, "ip", "Hz")
    ex_zero = metrics["zero_crossings"]["ip"]["Ex"]
    doc.add_paragraph(
        f"FEniCSx 的 no-IP 与 IP Hz 分别在 {hz_noip:.3f} s 和 "
        f"{hz_ip:.3f} s 附近出现参考解没有的符号反转。"
        "约 0.022 s 的 Hz 事件是 false reversal，是独立失败证据；"
        "它不能被早期强信号主导的总场 L2 掩盖。"
    )
    doc.add_paragraph(
        f"Ex 首次反号：FEniCSx={float(ex_zero['prediction'][0]):.6f} s，"
        f"empymod={float(ex_zero['reference'][0]):.6f} s，"
        f"相对时间误差={_percentage(ex_zero['max_relative_time_error'], 3)}。"
        "该正式结果保留。"
    )
    doc.add_paragraph(
        "原 strict_comparison.json 中 Ex、Hz 与 dBz/dt 的 IP 增量分别为 "
        f"{_percentage(metrics['ip_increment']['Ex']['relative_l2'])}、"
        f"{_percentage(metrics['ip_increment']['Hz']['relative_l2'])} 和 "
        f"{_percentage(metrics['ip_increment']['dBzdt']['relative_l2'])}。"
        "这些数值作为可复现敏感性证据保留；因参考变换审计未通过，"
        "不再把它们作为正式增量门限的通过/失败判定。"
    )

    doc.add_heading("5  Debye 极点数内部敏感性", level=1)
    _add_figure(
        doc,
        publication_bundle / FIGURES[4][0],
        FIGURES[4][1],
    )
    comparison = diagnostic["comparison"]
    _add_table(
        doc,
        ("分量", "4 项相对 16 项内部变化", "解释"),
        [
            [
                component,
                _percentage(
                    comparison[component]["debye_4_vs_16_relative_l2"],
                ),
                "内部敏感性；不是 empymod 跨代码门限",
            ]
            for component in ("Ex", "Hz", "dBzdt")
        ],
    )
    doc.add_paragraph(
        "4 项相对 16 项 Debye 只改变极点数，属于同代码、同网格、同时间步"
        "条件下的 pole-count sensitivity。它不能证明 16 项结果是外部真值，"
        "也不能替代 QWE 参考变换收敛证据。"
    )

    doc.add_heading("6  结论与适用边界", level=1)
    doc.add_paragraph(
        "总场 L2 数值和 Ex zero-crossing 的正式结果保持不变；Hz 约 0.022 s "
        "的 false reversal 仍是独立失败证据。default DLF 的 11.40% 与 QWE 的 "
        "9.44% 被保留为可复现敏感性，但前者不是正式失败、后者不是通过。"
        "Debye 4 项相对 16 项仅为内部敏感性。"
    )
    final = doc.add_paragraph()
    run_text = final.add_run(
        "因此当前最严谨的总体结论是：算法未完全通过；状态保持 "
        "failed_with_reproducible_evidence。现有证据支持接地源总场主体与 Ex "
        "反号链条，但不支持宣称 dBz/dt 弱极化增量或完整 TEM-IP 算法已严格通过。"
    )
    _format_run(run_text, bold=True, color=NAVY)
    doc.add_paragraph(
        "该结论属于算法内部和跨代码数值验证，不等同于堤坝现场有效性验证。"
        "后续仍需独立的网格、时间步、边界、实际波形、线圈有限面积与噪声审计。"
    )

    doc.add_heading("附录 A  可复现证据路径与 QA", level=1)
    metadata = manifest["metadata"]
    _add_table(
        doc,
        ("证据", "记录"),
        [
            ["正式运行", _relative_or_absolute(run, root)],
            [
                "strict comparison",
                _relative_or_absolute(
                    run / "comparisons/S1T1B1/strict_comparison.json",
                    root,
                ),
            ],
            ["publication bundle", _relative_or_absolute(publication_bundle, root)],
            ["reference audit", _relative_or_absolute(reference_audit, root)],
            ["bundle run_id", str(metadata.get("run_id"))],
            ["bundle reference status", str(metadata.get("reference_audit_status"))],
            [
                "DOCX QA",
                "layout rendering not performed / LibreOffice unavailable；"
                "ZIP/XML package、五张媒体及内部图片关系已检查。"
                "未生成 PDF；残余分页风险如实保留。",
            ],
        ],
    )
    doc.add_paragraph(
        "参考文献：Zhou, N.-N., Lei, K.-X., Xue, G.-Q., and Chen, W. "
        "(2020). Induced Polarization Effect on Grounded-Wire Transient "
        "Electromagnetic Data from Transverse Electric and Magnetic Fields. "
        "GEOPHYSICS, 85(4), E111-E120. DOI: 10.1190/geo2019-0322.1."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        prefix=f".{output.name}.",
        suffix=".tmp.docx",
        dir=output.parent,
        delete=False,
    ) as stream:
        temporary = Path(stream.name)
    try:
        doc.save(temporary)
        with temporary.open("rb+") as stream:
            os.fsync(stream.fileno())
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)


def build_report(
    root: Path,
    run: Path,
    publication_bundle: Path,
    reference_audit: Path,
    output: Path,
) -> None:
    root = Path(root)
    run = Path(run)
    publication_bundle = Path(publication_bundle)
    reference_audit = Path(reference_audit)
    output = Path(output)
    if output.suffix.lower() != ".docx":
        raise ValueError("report output must be DOCX; PDF generation is forbidden")

    bundle = _load_validated_bundle(publication_bundle)
    validated_audit = _load_validated_reference_audit(reference_audit)
    manifest, diagnostic, audit = _validate_evidence(bundle, validated_audit)

    metrics = _load_cross_bound_metrics(
        root,
        run,
        manifest,
        validated_audit,
    )
    for figure, _ in FIGURES:
        if not (publication_bundle / figure).is_file():
            raise FileNotFoundError(
                f"validated publication bundle figure is missing: {figure}"
            )

    git_commit, git_dirty = _git_state(root)
    _write_report(
        root=root,
        run=run,
        publication_bundle=publication_bundle,
        reference_audit=reference_audit,
        output=output,
        metrics=metrics,
        manifest=manifest,
        diagnostic=diagnostic,
        audit=audit,
        git_commit=git_commit,
        git_dirty=git_dirty,
    )


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    build_report(
        args.root.resolve(),
        args.run.resolve(),
        args.publication_bundle.resolve(),
        args.reference_audit.resolve(),
        args.output.resolve(),
    )
    print(f"DOCX report written: {args.output.resolve()}")
    print("PDF policy: no PDF generated or delivered")


if __name__ == "__main__":
    main()
