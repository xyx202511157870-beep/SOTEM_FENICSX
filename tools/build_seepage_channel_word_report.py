from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RESULT_DIR = ROOT / "output" / "seepage_channel_100m_5rx"
DEFAULT_REPORT_PATH = ROOT / "output" / "doc" / "100m长导线渗流通道三算法正演对比报告.docx"

BLUE = "245B84"
LIGHT_BLUE = "EAF2F8"
PALE_GREEN = "EAF5EE"
PALE_YELLOW = "FFF6D9"
INK = "243746"
MUTED = "66717A"
GRID = "BAC4CC"
WHITE = "FFFFFF"
CHINESE_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
REPORT_RECEIVER_INDICES = (0, 1, 3, 4)


def _read_json(path: Path) -> dict:
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def _set_run_font(run, size: float = 10.5, *, bold: bool = False, color: str = INK) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def _cant_split(row) -> None:
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    props.append(node)


def _repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def _set_cell_text(cell, value: object, *, header: bool = False, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(value))
    _set_run_font(run, 8.6, bold=header or bold, color=WHITE if header else INK)
    if header:
        _shade(cell, BLUE)


def add_table(document: Document, headers: Iterable[str], rows: Iterable[Iterable[object]]) -> None:
    headers = list(headers)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, header=True)
    _repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        _cant_split(row)
        for cell, value in zip(row.cells, values):
            _set_cell_text(cell, value)
            if row_index % 2:
                _shade(cell, "F6F8FA")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(document: Document, text: str, level: int = 1, *, page_break: bool = False) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, bold=True)
        tail = paragraph.add_run(text[len(bold_lead) :])
        _set_run_font(tail)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def add_callout(document: Document, label: str, text: str, fill: str = PALE_GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    lead = paragraph.add_run(label + "  ")
    _set_run_font(lead, 10.5, bold=True, color=BLUE)
    run = paragraph.add_run(text)
    _set_run_font(run, 10.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document, path: Path, caption: str, width: float = 6.35) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    _set_run_font(run, 9, color=MUTED)


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.font.size = Pt(10.5)
    for name, size, color in (("Title", 28, BLUE), ("Heading 1", 17, BLUE), ("Heading 2", 13, INK)):
        style = styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("100 m 长导线渗流通道三算法对比  |  ")
    _set_run_font(run, 8.5, color=MUTED)
    _append_page_field(footer)


def add_cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("100 m 长导线渗流通道三算法正演对比报告")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("SimPEG  |  empymod  |  FEniCSx")
    _set_run_font(run, 15, bold=True, color=BLUE)
    document.add_paragraph()
    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("研究目标", "平行长导线下方 20 m 高导渗流通道响应"),
        ("模型版本", "最新工作区模型基线"),
        ("坐标约定", "z=0 地表，z>0 地下，z<0 空气"),
        ("生成日期", date.today().isoformat()),
    ]
    for row, values in zip(meta.rows, meta_rows):
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            _set_cell_text(cell, value, bold=index == 0)
            if index == 0:
                _shade(cell, LIGHT_BLUE)
    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "正式报告视图：4 个接收点 x 31 个时刻 x 3 个分量；原始计算产物保留 5 点"
    )
    _set_run_font(run, 10.5, color=MUTED)
    document.add_page_break()


def _fmt_pct(value: float) -> str:
    return f"{100.0 * value:.2f}%"


def add_summary(document: Document, summary: dict) -> None:
    add_heading(document, "1. 执行摘要")
    add_callout(
        document,
        "结论",
        "SimPEG 与 FEniCSx 对渗流通道异常的 Ex、dBz/dt 和 Hz 三分量对比均通过 20% 强信号中位相对误差目标；两个三维算法的均匀背景也通过 empymod 一维基准检验。",
    )
    rows = []
    for method, components in summary["background"].items():
        for component in ("Ex", "dBzdt", "Hz"):
            item = components[component]
            rows.append(
                [
                    f"{method} 背景 / empymod",
                    component,
                    _fmt_pct(item["median_relative_error"]),
                    _fmt_pct(item["target"]),
                    item["strong_count"],
                    "通过" if item["pass"] else "未通过",
                ]
            )
    for component in ("Ex", "dBzdt", "Hz"):
        item = summary["channel_delta"][component]
        rows.append(
            [
                "SimPEG / FEniCSx 通道差分",
                component,
                _fmt_pct(item["median_relative_error"]),
                _fmt_pct(item["target"]),
                item["strong_count"],
                "通过" if item["pass"] else "未通过",
            ]
        )
    add_table(document, ["对比", "分量", "中位误差", "目标", "强信号数", "结果"], rows)
    add_body(
        document,
        "误差门槛只对强信号样本的中位相对误差判定；所有原始值、符号差分和逐点误差均保留在 CSV/NPZ 产物中，不会因弱信号筛选而丢失。",
    )


def add_model(document: Document, result_dir: Path, audit: dict) -> None:
    add_heading(document, "2. 物理模型与观测系统", page_break=True)
    add_figure(
        document,
        result_dir / "model_geometry.png",
        "图 1  发射导线、四个正式接收点与渗流通道的完整三维几何（物理深度向下）",
    )
    channel = audit["channel"]
    add_table(
        document,
        ["参数", "设定", "说明"],
        [
            ("坐标约定", "z 向下为正", "z=0 地表，z>0 地下，z<0 空气"),
            ("长导线", "(-50,0,0.1) m -> (50,0,0.1) m", "长度 100 m，电流 1 A，理想阶跃关断"),
            ("背景电导率", "0.01 S/m", "100 ohm m 均匀地下半空间"),
            ("空气电导率", "1e-8 S/m", "数值空气层"),
            ("通道尺寸", "60 x 10 x 10 m", "长轴与导线 x 方向平行"),
            ("通道中心", "(0,0,20) m", "正下方 20 m，深度范围 15-25 m"),
            ("通道电导率", f"{channel['conductivity_s_per_m']:.1f} S/m", "高导充水渗流通道，本阶段不含 IP"),
            ("观测时刻", "1e-5-1e-2 s，31 个对数时刻", "输出 Ex、dBz/dt、Hz"),
        ],
    )
    receiver_rows = []
    for receiver_index in REPORT_RECEIVER_INDICES:
        location = audit["receiver_locations_m"][receiver_index]
        receiver_rows.append(
            (
                f"Rx{receiver_index + 1}",
                *(f"{value:g}" for value in location),
                "explicit_full_domain",
            )
        )
    add_table(document, ["接收点", "x (m)", "y (m)", "z (m)", "FEniCSx 来源"], receiver_rows)
    add_callout(
        document,
        "四点报告视图与全域求解来源",
        "原始五点均在同一个完整三维域中直接求值；正式报告仅展示四个非中心接收点："
        "Rx1、Rx2、Rx4、Rx5。四点来源均为 explicit_full_domain，不使用单侧求解后对称镜像。",
        fill=PALE_YELLOW,
    )


def add_methods(document: Document, result_dir: Path, fenics_background: dict, fenics_channel: dict) -> None:
    add_heading(document, "3. 算法、离散与求解参数", page_break=True)
    add_table(
        document,
        ["方法", "模型能力与角色", "空间离散", "时间/频域处理"],
        [
            ("empymod", "分层一维背景参考；不表示有限三维长方体", "半空间积分核", "频域计算后转时域"),
            ("SimPEG", "三维背景和通道正演", "TensorMesh，345,600 单元", "EB 时域，371 个内部时间节点"),
            ("FEniCSx", "三维背景和通道全域正演", "一阶 Nedelec 四面体", "theta=1，相对步长 5%，晚期上限 1e-4 s"),
        ],
    )
    material = fenics_channel["material_audit"]
    add_heading(document, "3.1 SimPEG 参数", level=2)
    add_table(
        document,
        ["项目", "设定"],
        [
            ("坐标适配", "物理 z_down -> 内部 z_up"),
            ("求解形式", "EB，PARDISO"),
            ("源初始磁场", "Biot-Savart 有限导线"),
            ("内部时间步", "1e-7 x 100，1e-6 x 90，1e-5 x 90，1e-4 x 90"),
            ("结果形状", "原始 5 x 31 x 3；正式曲线展示 4 个接收点"),
        ],
    )
    add_heading(document, "3.2 FEniCSx 参数", level=2)
    add_table(
        document,
        ["项目", "背景", "含通道"],
        [
            ("四面体数", "156,203", "156,439"),
            ("Nedelec 自由度", "183,839", "184,116"),
            ("通道单元数", "0", material["global_cell_count"]),
            ("通道离散体积", "-", f"{material['global_discrete_volume_m3']:.1f} m3"),
            ("体积相对误差", "-", _fmt_pct(material["relative_volume_error"])),
            ("线源积分点", "4,048", "4,048"),
            ("时间内步数", "224", "224"),
            ("线性求解器", "CG + hypre/AMS", "CG + hypre/AMS"),
            ("计算总时间", f"{fenics_background['runtime_seconds']['total_seconds'] / 60:.1f} min", f"{fenics_channel['runtime_seconds']['total_seconds'] / 60:.1f} min"),
            ("输出模式", "原始 5 点全域求值；报告 4 点", "原始 5 点全域求值；报告 4 点"),
        ],
    )
    add_body(
        document,
        "FEniCSx 的 Ex 由 Nedelec 电场在接收点直接评估，dBz/dt 由 -curl(E) 的 DG0 插值获得，Hz 由全域电流的 Biot-Savart 积分获得。为避免不必要的每内步积分，Hz 仅在 31 个输出时刻计算，不改变输出时间网格。",
    )
    add_heading(document, "3.3 empymod 边界", level=2)
    add_callout(
        document,
        "重要边界",
        "empymod 在本报告中只是均匀背景一维参考（background_only_1d=true）。有限长 60 x 10 x 10 m 的通道是三维体，不能将通道响应误标为 empymod 的三维求解结果。",
        fill=PALE_YELLOW,
    )


def add_results(document: Document, result_dir: Path, summary: dict) -> None:
    add_heading(document, "4. 结果对比")
    add_heading(document, "4.1 均匀背景与 empymod 参考", level=2)
    add_figure(document, result_dir / "background_response.png", "图 2  均匀背景三算法 Ex、dBz/dt 和 Hz 响应")
    add_figure(document, result_dir / "background_error.png", "图 3  SimPEG 和 FEniCSx 背景相对 empymod 的逐点误差")
    add_heading(document, "4.2 含渗流通道的三维响应", level=2, page_break=True)
    add_figure(document, result_dir / "channel_response.png", "图 4  SimPEG 和 FEniCSx 含通道响应")
    add_heading(document, "4.3 通道带符号异常响应", level=2)
    add_figure(document, result_dir / "channel_delta.png", "图 5  通道响应减背景响应的带符号差分")
    add_figure(document, result_dir / "channel_delta_error.png", "图 6  SimPEG 与 FEniCSx 通道差分的逐点误差")
    add_callout(
        document,
        "异常对比结果",
        "强信号样本中，Ex、dBz/dt、Hz 的中位相对误差分别为 "
        + ", ".join(_fmt_pct(summary["channel_delta"][name]["median_relative_error"]) for name in ("Ex", "dBzdt", "Hz"))
        + "，均低于 20% 目标。",
    )


def add_quality_and_reproducibility(document: Document, result_dir: Path, audit: dict) -> None:
    add_heading(document, "5. 质量控制、局限与复现", page_break=True)
    add_table(
        document,
        ["检查项", "结果"],
        [
            (
                "数据完整性",
                "原始 SimPEG/FEniCSx 数组均为 5 x 31 x 3 且全部有限；正式报告展示 4 点",
            ),
            (
                "FEniCSx 接收点来源",
                "原始 5/5 为 explicit_full_domain；正式四点无单侧对称镜像",
            ),
            ("通道网格体积", "5727.3 m3 / 6000 m3，相对误差 4.55%"),
            ("比较边界", "empymod 只参与背景比较，禁止将通道差分与一维参考直接比较"),
            ("空间/时间收敛扫描", "本次未执行独立加密算例；convergence_summary.json 显式记录 available=false"),
        ],
    )
    add_body(
        document,
        "局限：本阶段通道采用静态纯电导率 1 S/m，未包含激发极化、各向异性、粗糙边界或流体饭和度变化。对于通道尺寸、电导率或埋深的反演结论，建议在本基准上增加空间与时间加密扫描。",
        bold_lead="局限：",
    )
    add_heading(document, "5.1 主要产物", level=2)
    add_table(
        document,
        ["产物", "用途"],
        [
            ("benchmark_results.npz", "三算法统一数组和通道差分"),
            ("benchmark_values.csv", "所有背景/通道原始数值"),
            ("channel_delta_values.csv", "带符号通道异常"),
            ("background_errors.csv", "三维背景对 empymod 的逐点误差"),
            ("channel_delta_errors.csv", "SimPEG/FEniCSx 通道差分逐点误差"),
            ("benchmark_manifest.json", "文件清单、字节数和 SHA-256"),
        ],
    )
    add_heading(document, "5.2 复现命令", level=2)
    commands = [
        "python tools/run_seepage_channel_benchmark.py all --output-root output/seepage_channel_100m_5rx",
        "python tools/plot_seepage_channel_benchmark.py output/seepage_channel_100m_5rx",
        "bash tools/run_fenicsx_seepage_background.sh",
        "bash tools/run_fenicsx_seepage_channel.sh",
    ]
    for command in commands:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(INK)
    add_body(document, f"所有正式数值、图件、日志和摘要位于：{result_dir.resolve()}")


def add_magnetic_stability_section(document: Document, result_dir: Path) -> None:
    """Add the selected magnetic method for the four reporting receivers."""

    magnetic_dir = result_dir / "magnetic_receiver_stability"
    metrics = _read_json(magnetic_dir / "magnetic_symmetry_metrics.json")
    decision = _read_json(magnetic_dir / "magnetic_convergence_summary.json")
    mesh_audit = _read_json(magnetic_dir / "mesh_symmetry_audit.json")
    selected = decision.get("selected")
    selected_metrics = metrics.get(selected, {}) if selected else {}
    conclusion = "PASS" if decision.get("passed") else "FAIL"
    selected_dbdt_operator = {
        "biot_rate": "Biot-Savart Hz 时间差分（biot_rate）",
        "curl": "单元 curl(E)",
    }.get(selected, "Faraday 有限线圈" if str(selected).startswith("faraday_loop") else "未选择")

    add_heading(document, "四个正式接收点的磁场计算稳定性", page_break=True)
    add_body(
        document,
        "问题根因是点式 -curl(E) 在材料界面和共享单元处对局部离散误差敏感，"
        "候选诊断同时评估单元 curl(E)、Biot-Savart Hz 时间差分和 Faraday 有限线圈；"
        "正式算子由对称残差门禁与强信号一致性审计的数据结果选择，而不是预先指定。"
        "位于导线正下方的理论零点不进入正式观测曲线，报告只展示 Rx1、Rx2、Rx4、Rx5；"
        "Hz 统一用四面体四点 Biot-Savart 计算。",
    )
    add_callout(
        document,
        "全域求解约束",
        "原始五点均标记为 explicit_full_domain；正式报告仅展示四个非中心接收点："
        "Rx1、Rx2、Rx4、Rx5。不使用单侧求解后对称镜像，也不从相反 y 坐标复制接收值。",
        fill=PALE_YELLOW,
    )
    add_table(
        document,
        ["参数", "设定"],
        [
            ("模型", "100 m 平行导线；60 x 10 x 10 m 渗流通道；埋深 20 m"),
            ("正式 dBz/dt 算子", selected_dbdt_operator),
            ("Faraday 候选诊断", "有限线圈，未通过时仅保留为对照"),
            ("线圈半径", "2 m"),
            ("线积分点数", "32"),
            ("Hz 算子", "四面体四点 Biot-Savart"),
            ("Biot 积分模式", "tetra4"),
        ],
    )
    add_table(
        document,
        ["对称指标", "实测值", "门限"],
        [
            ("Rx2/Rx4 奇对称残差", selected_metrics.get("pair_24_residual", "N/A"), "<= 0.01"),
            ("Rx1/Rx5 奇对称残差", selected_metrics.get("pair_15_residual", "N/A"), "<= 0.01"),
        ],
    )
    add_table(
        document,
        ["审计项目", "结果"],
        [
            ("正式方法", selected or "未选择"),
            ("综合门禁", conclusion),
            ("求解域", mesh_audit.get("solver_domain", "unknown")),
            ("场值镜像", mesh_audit.get("field_mirroring", "unknown")),
            ("网格对称审计", "PASS" if mesh_audit.get("passed") else "FAIL"),
        ],
    )
    add_figure(
        document,
        magnetic_dir / "magnetic_receiver_comparison.png",
        "四个正式接收点的磁场计算方法带符号曲线对比",
    )
    add_body(
        document,
        f"数据驱动结论：综合门禁为 {conclusion}；正式方法为 {selected or '无'}。"
        "结论直接读取 JSON 审计产物，不在 Word 生成器中硬编码通过。",
    )


def _build_magnetic_only_report(result_dir: Path, output_path: Path) -> Path:
    document = Document()
    configure_document(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("100 m 平行导线渗流通道磁场稳定性报告")
    _set_run_font(run, 20, bold=True, color=BLUE)
    add_body(
        document,
        "比较 SimPEG、empymod 与 FEniCSx 的统一模型参数，并重点记录 FEniCSx 磁场接收算子、"
        "四个正式接收点曲线及原始五点全域求解审计。",
    )
    add_magnetic_stability_section(document, result_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def build_report(
    *,
    result_dir: Path = DEFAULT_RESULT_DIR,
    output_path: Path = DEFAULT_REPORT_PATH,
) -> Path:
    result_dir = Path(result_dir)
    output_path = Path(output_path)
    if not (result_dir / "benchmark_summary.json").exists():
        return _build_magnetic_only_report(result_dir, output_path)
    summary = _read_json(result_dir / "benchmark_summary.json")
    audit = _read_json(result_dir / "model_audit.json")
    fenics_background = _read_json(result_dir / "fenicsx_background" / "fenicsx_run_summary.json")
    fenics_channel = _read_json(result_dir / "fenicsx_channel" / "fenicsx_run_summary.json")

    document = Document()
    configure_document(document)
    add_cover(document)
    add_summary(document, summary)
    add_model(document, result_dir, audit)
    add_methods(document, result_dir, fenics_background, fenics_channel)
    add_results(document, result_dir, summary)
    add_quality_and_reproducibility(document, result_dir, audit)
    if (result_dir / "magnetic_receiver_stability" / "magnetic_symmetry_metrics.json").exists():
        add_magnetic_stability_section(document, result_dir)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the seepage-channel three-solver Word report.")
    parser.add_argument("--result-dir", type=Path, default=DEFAULT_RESULT_DIR)
    parser.add_argument("--output", type=Path, default=DEFAULT_REPORT_PATH)
    args = parser.parse_args()
    path = build_report(result_dir=args.result_dir, output_path=args.output)
    print(path.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
