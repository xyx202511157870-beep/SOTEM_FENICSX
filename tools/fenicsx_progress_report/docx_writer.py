"""Render the curated FEniCSx progress archive as a compact Word document."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Any, Mapping, Sequence

from PIL import Image
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
PAGE_MARGIN_MM = 17
CONTENT_WIDTH_DXA = int(Mm(PAGE_WIDTH_MM - 2 * PAGE_MARGIN_MM).twips)

BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
TEXT_COLOR = "203A49"
MUTED_COLOR = "60747F"
ACCENT_COLOR = "287CA3"
DEEP_ACCENT_COLOR = "1F4D78"
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER_COLOR = "A7B5BD"

DIRECTIVE_RE = re.compile(r"^\[\[([A-Z_]+)(?::([^\]]+))?\]\]$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
ORDERED_LIST_RE = re.compile(r"^\d+\.\s+(.+)$")


class UnknownFigureError(KeyError):
    """Raised when the source names a figure absent from the manifest."""


class UnknownDirectiveError(ValueError):
    """Raised when the source contains a directive the renderer cannot handle."""


@dataclass(frozen=True)
class _ImageSize:
    width: Mm
    height: Mm


def _set_run_font(run: Any, name: str, size: float, color: str = TEXT_COLOR) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_style_font(style: Any, name: str, size: float, color: str) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _configure_styles(document: DocumentObject) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, BODY_FONT, 10, TEXT_COLOR)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in (
        ("Title", 24, DEEP_ACCENT_COLOR, 0, 10),
        ("Heading 1", 16, ACCENT_COLOR, 10, 6),
        ("Heading 2", 13, ACCENT_COLOR, 8, 4),
        ("Heading 3", 11.5, DEEP_ACCENT_COLOR, 6, 3),
        ("Heading 4", 10.5, DEEP_ACCENT_COLOR, 5, 2),
    ):
        style = styles[name]
        _set_style_font(style, BODY_FONT, size, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    _set_style_font(caption, BODY_FONT, 8.25, MUTED_COLOR)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        _set_style_font(style, BODY_FONT, 9.75, TEXT_COLOR)
        style.paragraph_format.left_indent = Mm(5)
        style.paragraph_format.first_line_indent = Mm(-2.5)
        style.paragraph_format.space_after = Pt(2)


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("第 ")
    _set_run_font(prefix, BODY_FONT, 8, MUTED_COLOR)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    suffix = paragraph.add_run(" 页")
    _set_run_font(suffix, BODY_FONT, 8, MUTED_COLOR)


def _configure_section(section: Any) -> None:
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(PAGE_MARGIN_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_MM)
    section.left_margin = Mm(PAGE_MARGIN_MM)
    section.right_margin = Mm(PAGE_MARGIN_MM)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    header = section.header.paragraphs[0]
    header.text = "FEniCSx 算法建立与调试进展"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(header.runs[0], BODY_FONT, 8, MUTED_COLOR)
    header.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    _add_page_number(footer)
    footer.paragraph_format.space_after = Pt(0)


def _set_cell_margins(cell: Any, *, top: int = 70, start: int = 100,
                      bottom: int = 70, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _column_widths(column_count: int, weights: Sequence[float] | None = None) -> list[int]:
    if not weights:
        weights = [1.0] * column_count
    total_weight = sum(weights)
    widths = [int(CONTENT_WIDTH_DXA * weight / total_weight) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = widths[index]
            _set_cell_margins(cell)


def _set_table_borders(table: Any, color: str = TABLE_BORDER_COLOR) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _add_inline_text(paragraph: Any, text: str, *, size: float | None = None,
                     color: str = TEXT_COLOR) -> None:
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            _set_run_font(run, BODY_FONT, size or 10, color)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, CODE_FONT, (size or 10) - 0.5, DEEP_ACCENT_COLOR)
        else:
            run = paragraph.add_run(token)
            _set_run_font(run, BODY_FONT, size or 10, color)


def _add_text_paragraph(document: DocumentObject, text: str, style: str | None = None) -> Any:
    paragraph = document.add_paragraph(style=style)
    _add_inline_text(paragraph, text)
    paragraph.paragraph_format.widow_control = True
    return paragraph


def _strip_table_cell(value: str) -> str:
    return value.strip().replace("\\|", "|")


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [_strip_table_cell(value) for value in re.split(r"(?<!\\)\|", stripped)]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _add_markdown_table(document: DocumentObject, lines: Sequence[str]) -> None:
    rows = [_split_table_row(line) for line in lines]
    if len(rows) >= 2 and _is_table_separator(lines[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    rows = [row + [""] * (column_count - len(row)) for row in rows]
    weights = []
    for column in range(column_count):
        longest = max(len(row[column]) for row in rows)
        weights.append(max(1.0, min(3.0, longest / 18.0)))
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_geometry(table, _column_widths(column_count, weights))
    _set_table_borders(table)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _add_inline_text(paragraph, value, size=8.4)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                _set_cell_shading(cell, TABLE_HEADER_FILL)
        if row_index == 0:
            _set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _resolve_image_path(project_root: Path, value: str) -> Path:
    path = Path(value)
    if not path.is_absolute():
        path = project_root / path
    if not path.is_file():
        raise FileNotFoundError(f"figure image not found: {path}")
    return path


def _validate_figure(record: Mapping[str, Any]) -> None:
    for field in ("id", "path", "caption", "status", "source_run", "problem", "change", "result"):
        if not str(record.get(field, "")).strip():
            raise ValueError(f"figure {record.get('id', '<unknown>')} requires {field}")


def _fit_image(path: Path, max_width_mm: float, max_height_mm: float) -> _ImageSize:
    with Image.open(path) as image:
        width, height = image.size
    if width <= 0 or height <= 0:
        raise ValueError(f"invalid image dimensions: {path}")
    scale = min(max_width_mm / width, max_height_mm / height)
    return _ImageSize(Mm(width * scale), Mm(height * scale))


def _set_drawing_alt_text(shape: Any, description: str) -> None:
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", description[:1000])
    doc_pr.set("title", description[:255])


def _add_body_figure(document: DocumentObject, record: Mapping[str, Any],
                     project_root: Path) -> None:
    _validate_figure(record)
    path = _resolve_image_path(project_root, str(record["path"]))
    dimensions = _fit_image(path, 176, 155)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    shape = paragraph.add_run().add_picture(
        str(path), width=dimensions.width, height=dimensions.height
    )
    _set_drawing_alt_text(shape, str(record["caption"]))

    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = caption.add_run(f"[{record['status']}] {record['caption']}")
    lead.bold = True
    _set_run_font(lead, BODY_FONT, 8.25, DEEP_ACCENT_COLOR)
    caption.add_run("\n")
    details = (
        f"来源：{record['source_run']} | 问题：{record['problem']} "
        f"| 修改：{record['change']} | 结果：{record['result']}"
    )
    _add_inline_text(caption, details, size=7.7, color=MUTED_COLOR)


def _add_atlas_cell(cell: Any, record: Mapping[str, Any], project_root: Path) -> None:
    _validate_figure(record)
    path = _resolve_image_path(project_root, str(record["path"]))
    dimensions = _fit_image(path, 78, 46)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_margins(cell, top=80, start=90, bottom=70, end=90)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    shape = paragraph.add_run().add_picture(
        str(path), width=dimensions.width, height=dimensions.height
    )
    _set_drawing_alt_text(shape, str(record["caption"]))

    caption = cell.add_paragraph()
    caption.paragraph_format.space_after = Pt(1)
    caption.paragraph_format.line_spacing = 1.0
    lead = caption.add_run(f"[{record['status']}] {record['caption']}")
    lead.bold = True
    _set_run_font(lead, BODY_FONT, 6.8, DEEP_ACCENT_COLOR)

    source = cell.add_paragraph()
    source.paragraph_format.space_after = Pt(1)
    source.paragraph_format.line_spacing = 1.0
    _add_inline_text(source, f"来源：{record['path']}", size=6.1, color=MUTED_COLOR)

    evidence = cell.add_paragraph()
    evidence.paragraph_format.space_after = Pt(0)
    evidence.paragraph_format.line_spacing = 1.0
    _add_inline_text(
        evidence,
        f"问题：{record['problem']} 修改：{record['change']} 结果：{record['result']}",
        size=6.1,
        color=TEXT_COLOR,
    )


def _add_figure_atlas(document: DocumentObject, figures: Sequence[Mapping[str, Any]],
                      project_root: Path) -> None:
    appendix = [figure for figure in figures if figure.get("placement") == "appendix"]
    if not appendix:
        _add_text_paragraph(document, "本次清单中没有附录图件。")
        return
    for start in range(0, len(appendix), 4):
        group = appendix[start:start + 4]
        table = document.add_table(rows=2, cols=2)
        _set_table_geometry(table, _column_widths(2))
        _set_table_borders(table, color="C7D1D7")
        for index, record in enumerate(group):
            row, column = divmod(index, 2)
            _add_atlas_cell(table.cell(row, column), record, project_root)
        if start + 4 < len(appendix):
            document.add_page_break()


def _add_git_timeline(document: DocumentObject,
                      timeline: Sequence[Mapping[str, Any]] | None) -> None:
    if not timeline:
        _add_text_paragraph(document, "未提供 Git 时间线数据。")
        return
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _set_table_geometry(table, _column_widths(3, [1.15, 0.75, 4.1]))
    _set_table_borders(table)
    headers = ("日期", "提交", "说明")
    for index, value in enumerate(headers):
        cell = table.cell(0, index)
        _set_cell_shading(cell, TABLE_HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _add_inline_text(paragraph, value, size=7.5)
        for run in paragraph.runs:
            run.bold = True
    _set_repeat_table_header(table.rows[0])
    for event in timeline:
        row = table.add_row()
        values = (
            str(event.get("date") or event.get("authored_at") or ""),
            str(event.get("short_hash") or event.get("hash") or "")[:10],
            str(event.get("subject") or event.get("message") or ""),
        )
        for index, value in enumerate(values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _add_inline_text(paragraph, value, size=7.1)
    _set_table_geometry(table, _column_widths(3, [1.15, 0.75, 4.1]))


def _validate_directives(source: str, figures_by_id: Mapping[str, Mapping[str, Any]]) -> None:
    for raw_line in source.splitlines():
        line = raw_line.strip()
        if not (line.startswith("[[") and line.endswith("]]")):
            continue
        match = DIRECTIVE_RE.fullmatch(line)
        if match is None:
            raise UnknownDirectiveError(f"unknown directive: {line}")
        name, argument = match.groups()
        if name == "FIGURE":
            if not argument or argument not in figures_by_id:
                raise UnknownFigureError(argument or "<missing id>")
        elif name in {"FIGURE_ATLAS", "GIT_TIMELINE", "PAGEBREAK"}:
            if argument:
                raise UnknownDirectiveError(f"directive {name} does not accept an argument")
        else:
            raise UnknownDirectiveError(f"unknown directive: {name}")


def _source_parts(source: str) -> tuple[str, list[str], list[str]]:
    lines = source.splitlines()
    title = "FEniCSx 算法建立与调试进展报告"
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        lines = lines[1:]
    first_chapter = next((index for index, line in enumerate(lines) if line.startswith("## ")), len(lines))
    return title, lines[:first_chapter], lines[first_chapter:]


def _add_cover(document: DocumentObject, title: str, metadata: Sequence[str]) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    accent = document.add_paragraph()
    accent.paragraph_format.space_after = Pt(10)
    run = accent.add_run("FEniCSx / TDEM / VALIDATION ARCHIVE")
    run.bold = True
    _set_run_font(run, CODE_FONT, 9, ACCENT_COLOR)

    heading = document.add_paragraph(style="Title")
    heading.paragraph_format.keep_with_next = True
    _add_inline_text(heading, title, size=24, color=DEEP_ACCENT_COLOR)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "18")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), ACCENT_COLOR)
    bottom.append(border)
    rule._p.get_or_add_pPr().append(bottom)

    for line in metadata:
        stripped = line.strip()
        if not stripped:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        _add_inline_text(paragraph, stripped, size=10, color=MUTED_COLOR)
    document.add_page_break()


def _add_contents(document: DocumentObject, chapter_lines: Sequence[str]) -> None:
    document.add_heading("目录", level=1)
    for line in chapter_lines:
        if line.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(0)
            paragraph.paragraph_format.space_after = Pt(2)
            _add_inline_text(paragraph, line[3:].strip(), size=9.5, color=DEEP_ACCENT_COLOR)
        elif line.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(6)
            paragraph.paragraph_format.space_after = Pt(1)
            _add_inline_text(paragraph, line[4:].strip(), size=8.5, color=MUTED_COLOR)
    document.add_page_break()


def _render_chapters(document: DocumentObject, lines: Sequence[str],
                     figures: Sequence[Mapping[str, Any]], project_root: Path,
                     timeline: Sequence[Mapping[str, Any]] | None) -> None:
    figures_by_id = {str(figure["id"]): figure for figure in figures}
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [raw, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        if line.startswith("#### "):
            document.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            _add_text_paragraph(document, line[2:].strip(), style="List Bullet")
        elif (ordered := ORDERED_LIST_RE.fullmatch(line)) is not None:
            _add_text_paragraph(document, ordered.group(1), style="List Number")
        elif line.startswith("[[") and line.endswith("]]" ):
            match = DIRECTIVE_RE.fullmatch(line)
            if match is None:
                raise UnknownDirectiveError(f"unknown directive: {line}")
            name, argument = match.groups()
            if name == "FIGURE":
                _add_body_figure(document, figures_by_id[str(argument)], project_root)
            elif name == "FIGURE_ATLAS":
                _add_figure_atlas(document, figures, project_root)
            elif name == "GIT_TIMELINE":
                _add_git_timeline(document, timeline)
            elif name == "PAGEBREAK":
                document.add_page_break()
            else:
                raise UnknownDirectiveError(f"unknown directive: {name}")
        else:
            _add_text_paragraph(document, line)
        index += 1


def build_docx(source: str, selection: Mapping[str, Any], *, project_root: Path,
               output_path: Path, timeline: Sequence[Mapping[str, Any]] | None = None) -> Path:
    """Build a compact A4 report from Markdown source and a curated figure manifest."""
    figures = list(selection.get("figures", []))
    figure_ids = [str(figure.get("id", "")) for figure in figures]
    duplicates = sorted({figure_id for figure_id in figure_ids if figure_ids.count(figure_id) > 1})
    if duplicates:
        raise ValueError(f"duplicate figure id: {', '.join(duplicates)}")
    figures_by_id = {str(figure["id"]): figure for figure in figures}
    _validate_directives(source, figures_by_id)

    document = Document()
    _configure_styles(document)
    _configure_section(document.sections[0])
    title, metadata, chapters = _source_parts(source)
    _add_cover(document, title, metadata)
    _add_contents(document, chapters)
    _render_chapters(document, chapters, figures, Path(project_root), timeline)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = title
    document.core_properties.subject = "FEniCSx 时域电性源瞬变电磁法算法建立、调试与验证证据档案"
    document.core_properties.author = "FEniCSx TDEM 项目组"
    document.core_properties.keywords = "FEniCSx, TDEM, COMSOL, empymod, 验证, 收敛性"
    document.save(output_path)
    return output_path
