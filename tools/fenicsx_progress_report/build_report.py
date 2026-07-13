"""Audit inputs and build the reproducible FEniCSx progress Word report."""

from __future__ import annotations

import argparse
from dataclasses import asdict, dataclass
import hashlib
import json
import os
from pathlib import Path
import sys
from typing import Any, Mapping, Sequence
from uuid import uuid4

from docx import Document

from .docx_writer import build_docx
from .selection import audit_selection
from .source_audit import audit_source


class ReportAuditError(RuntimeError):
    """Raised when the evidence manifest or report source fails its audit."""


@dataclass(frozen=True)
class BuildSummary:
    output_path: str
    byte_size: int
    sha256: str
    figure_count: int
    body_figure_count: int
    appendix_figure_count: int
    excluded_figure_count: int
    git_event_count: int


def _load_json(path: Path) -> Any:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _format_selection_errors(report: Any) -> str:
    details = {
        name: value
        for name, value in report.__dict__.items()
        if value
    }
    return json.dumps(details, ensure_ascii=False, sort_keys=True)


def _format_source_errors(report: Any) -> str:
    return json.dumps(
        {
            "codes": list(report.codes),
            "missing_sections": list(getattr(report, "missing_sections", [])),
            "unknown_figure_ids": list(getattr(report, "unknown_figure_ids", [])),
            "missing_body_figure_ids": list(
                getattr(report, "missing_body_figure_ids", [])
            ),
        },
        ensure_ascii=False,
        sort_keys=True,
    )


def build_report(
    *,
    project_root: Path,
    source_path: Path,
    selection_path: Path,
    inventory_path: Path,
    timeline_path: Path,
    output_path: Path,
) -> BuildSummary:
    """Audit all inputs, render to a temporary DOCX, then atomically publish it."""
    project_root = Path(project_root).resolve()
    source = Path(source_path).read_text(encoding="utf-8")
    selection: Mapping[str, Any] = _load_json(Path(selection_path))
    inventory_payload: Mapping[str, Any] = _load_json(Path(inventory_path))
    timeline: Sequence[Mapping[str, Any]] = _load_json(Path(timeline_path))

    selection_report = audit_selection(
        list(inventory_payload.get("images", [])), selection
    )
    if not selection_report.ok:
        raise ReportAuditError(
            f"selection audit failed: {_format_selection_errors(selection_report)}"
        )

    source_report = audit_source(source, selection)
    if not source_report.ok:
        raise ReportAuditError(
            f"source audit failed: {_format_source_errors(source_report)}"
        )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.parent / f".{output_path.stem}.{uuid4().hex}.tmp.docx"
    try:
        build_docx(
            source,
            selection,
            project_root=project_root,
            output_path=temporary,
            timeline=timeline,
        )
        reopened = Document(temporary)
        if not reopened.paragraphs and not reopened.tables:
            raise RuntimeError("generated Word document is empty")
        os.replace(temporary, output_path)
    finally:
        if temporary.exists():
            temporary.unlink()

    figures = list(selection.get("figures", []))
    return BuildSummary(
        output_path=str(output_path.resolve()),
        byte_size=output_path.stat().st_size,
        sha256=_sha256(output_path),
        figure_count=len(figures),
        body_figure_count=sum(
            figure.get("placement") == "body" for figure in figures
        ),
        appendix_figure_count=sum(
            figure.get("placement") == "appendix" for figure in figures
        ),
        excluded_figure_count=len(selection.get("excluded", [])),
        git_event_count=len(timeline),
    )


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", required=True, type=Path)
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--selection", required=True, type=Path)
    parser.add_argument("--inventory", required=True, type=Path)
    parser.add_argument("--timeline", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--summary", type=Path)
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    summary = build_report(
        project_root=args.root,
        source_path=args.source,
        selection_path=args.selection,
        inventory_path=args.inventory,
        timeline_path=args.timeline,
        output_path=args.output,
    )
    payload = json.dumps(asdict(summary), ensure_ascii=False, indent=2) + "\n"
    if args.summary is not None:
        args.summary.parent.mkdir(parents=True, exist_ok=True)
        args.summary.write_text(payload, encoding="utf-8")
    print(payload, end="")
    return 0


if __name__ == "__main__":
    sys.exit(main())

