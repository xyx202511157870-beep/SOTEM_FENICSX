#!/usr/bin/env python3
"""Build the formal Chinese Word report from fully verified seepage artifacts."""

from __future__ import annotations

import argparse
from datetime import date
import json
from pathlib import Path
import sys
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import yaml


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from atem3d.seepage_verification import VerificationGateError  # noqa: E402
from tools.build_seepage_channel_word_report import (  # noqa: E402
    BLUE,
    CHINESE_FONT,
    INK,
    LATIN_FONT,
    MUTED,
    PALE_GREEN,
    PALE_YELLOW,
    _append_page_field,
    _set_run_font,
    add_body,
    add_callout,
    add_figure,
    add_heading,
    add_table,
)


FORMAL_RECEIVERS = (0, 1, 3, 4)
FIGURES = (
    ("verified_model_geometry.png", "图 1  物理深度向下的三维模型与四个正式观测点"),
    ("verified_background_decay.png", "图 2  均匀背景场的常规绝对值衰减曲线"),
    ("verified_channel_decay.png", "图 3  含渗流通道总场的常规绝对值衰减曲线"),
    ("verified_signed_anomaly.png", "图 4  通道减背景的带符号异常响应"),
    ("verified_relative_anomaly.png", "图 5  通道相对异常百分比"),
    ("verified_conductivity_sweep.png", "图 6  电导率扫描的异常能量趋势"),
    ("verified_volume_sweep.png", "图 7  异常体体积扫描的异常能量趋势"),
    ("verified_convergence.png", "图 8  空间与时间三层收敛结果"),
    ("verified_parity.png", "图 9  四个正式观测点的偶/奇对称残差"),
    ("verified_two_solver_anomaly.png", "图 10  SimPEG 与 FEniCSx 三维异常差分对比"),
)


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _metric_text(gate: dict[str, Any]) -> str:
    if "normalized_l2" in gate:
        values = ", ".join(f"{value:.2e}" for value in gate["normalized_l2"])
        threshold = gate.get("threshold")
        suffix = f"；阈值 {threshold:.1e}" if threshold is not None else ""
        return f"归一化 L2=[{values}]{suffix}"
    if "fine_medium" in gate:
        stats = gate["fine_medium"]
        return f"细-中：median={stats['median']:.1%}，P95={stats['p95']:.1%}"
    if "energies" in gate:
        values = " → ".join(f"{value:.3g}" for value in gate["energies"])
        return f"归一化异常能量：{values}"
    if "comparisons" in gate:
        medians = []
        for name, comparison in gate["comparisons"].items():
            component_values = [
                item["median"] for item in comparison.get("components", {}).values()
            ]
            if component_values:
                medians.append(f"{name} 最大分量中位差={max(component_values):.1%}")
        if medians:
            return "；".join(medians)
    if "components" in gate:
        medians = [
            item.get("median")
            for item in gate["components"].values()
            if item.get("median") is not None
        ]
        if medians:
            return f"最大分量中位差={max(medians):.1%}"
    if "median_threshold" in gate:
        return (
            f"median≤{gate['median_threshold']:.1%}，"
            f"P95≤{gate['p95_threshold']:.1%}"
        )
    if "selected" in gate:
        return f"正式算子={gate['selected']}"
    if "threshold" in gate:
        return f"阈值={gate['threshold']}"
    return "证据文件完整且指标满足门槛"


def _gate_status_rows(summary: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for name in summary["required_gates"]:
        gate = summary["gates"][name]
        rows.append(
            [
                name,
                _metric_text(gate),
                "通过" if gate.get("available") and gate.get("pass") else "未通过",
            ]
        )
    return rows


def _configure_document(document: Document, fingerprint: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.font.size = Pt(10.5)
    for name, size, color in (
        ("Title", 26, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, INK),
    ):
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"60×1×1 m 渗流通道科学验证 | {fingerprint[:12]}… | ")
    _set_run_font(run, 8.5, color=MUTED)
    _append_page_field(footer)


def _add_cover(document: Document, fingerprint: str) -> None:
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("100 m 平行导线下 60×1×1 m 渗流通道\n三维正演科学验证报告")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("SimPEG  |  FEniCSx  |  empymod")
    _set_run_font(run, 14, bold=True, color=BLUE)
    document.add_paragraph()
    add_table(
        document,
        ["项目", "内容"],
        [
            ("坐标约定", "z=0 为地表，z 向下为正"),
            ("通道埋深", "中心 20 m；深度范围 19.5–20.5 m"),
            ("正式观测点", "Rx1、Rx2、Rx4、Rx5；均为完整三维域直接求值"),
            ("模型 SHA-256", fingerprint),
            ("生成日期", date.today().isoformat()),
        ],
    )
    add_callout(
        document,
        "正式版本",
        "本报告仅在所有数值、物理、收敛、对称性和独立三维参考门槛全部通过后生成。",
        fill=PALE_GREEN,
    )
    document.add_page_break()


def _add_model_section(document: Document, result_dir: Path, audit: dict[str, Any]) -> None:
    add_heading(document, "1. 物理模型与观测系统")
    add_figure(document, result_dir / FIGURES[0][0], FIGURES[0][1])
    channel = audit["channel"]
    add_table(
        document,
        ["参数", "设定"],
        [
            ("发射导线", "(-50,0,0.1) m 至 (50,0,0.1) m；1 A；理想阶跃关断"),
            ("空气/背景电导率", "1e-8 / 0.01 S/m"),
            ("通道中心与尺寸", f"(0,0,20) m；{'×'.join(f'{v:g}' for v in channel['size_m'])} m"),
            ("通道电导率", f"{channel['conductivity_s_per_m']:g} S/m"),
            ("输出", "1e-5–1e-2 s 的 31 个对数时刻；Ex、dBz/dt、Hz"),
            ("正式观测曲线", "Rx1、Rx2、Rx4、Rx5；不采用单侧求解后镜像"),
        ],
    )


def _add_solver_section(document: Document, result_dir: Path) -> None:
    add_heading(document, "2. 求解器与离散参数", page_break=True)
    manifest = _read_json(result_dir / "verification_case_manifest.json")
    cases = manifest["cases"]
    unique = len({item["execution_fingerprint"] for item in cases})
    add_table(
        document,
        ["方法", "作用", "主要设定"],
        [
            ("empymod", "仅验证均匀半空间背景", "一维半空间半解析参考；不声称求解三维有限异常体"),
            ("SimPEG", "三维背景、异常体与扫描", "EB 时域有限体积；TensorMesh；PARDISO"),
            ("FEniCSx", "完整三维域独立有限元", "一阶 Nédélec 四面体；θ=1；CG+hypre/AMS"),
        ],
    )
    add_body(
        document,
        f"验证矩阵共 {len(cases)} 个逻辑算例、{unique} 个唯一执行指纹；重复的物理与数值配置只复用经过指纹核验的原始结果。",
    )
    add_table(
        document,
        ["控制变量", "水平"],
        [
            ("通道电导率", "0.01、0.02、0.1、1.0 S/m"),
            ("横截面", "1×1、2×2、10×10 m；长度固定 60 m"),
            ("局部空间尺度", "0.5、0.25、0.125 m"),
            ("内部时间步因子", "1、1/2、1/4"),
            ("FEniCSx 正式磁场", "dBz/dt=biot_rate；Hz=四面体四点 Biot-Savart"),
        ],
    )


def _add_gate_section(document: Document, summary: dict[str, Any]) -> None:
    add_heading(document, "3. 科学验证门槛", page_break=True)
    add_callout(
        document,
        "总体验收",
        f"全部 {len(summary['required_gates'])} 项强制门槛均通过；失败项为 0。",
    )
    add_table(document, ["门槛", "关键指标", "结论"], _gate_status_rows(summary))
    add_body(
        document,
        "正式误差统计仅使用四个非中心观测点和各分量的强信号时间窗。中心位置不绘制为观测曲线，也不参与跨算法误差百分位。",
    )


def _add_result_figures(document: Document, result_dir: Path) -> None:
    add_heading(document, "4. 结果与三算法对比", page_break=True)
    for filename, caption in FIGURES[1:]:
        add_figure(document, result_dir / filename, caption)
    add_body(
        document,
        "总场仍可呈现近似均匀半空间的单调衰减，因为一次背景场占主导。异常体证据由通道减背景的带符号差值、相对异常、参数扫描、收敛趋势以及 SimPEG 与完整域 FEniCSx 的独立三维差分对比共同给出。",
    )


def _add_reproducibility(document: Document, result_dir: Path, summary: dict[str, Any]) -> None:
    add_heading(document, "5. 可复现性、限制与文件", page_break=True)
    add_table(
        document,
        ["证据", "位置"],
        [
            ("最终门槛", str((result_dir / "verification_summary.json").resolve())),
            ("算例矩阵", str((result_dir / "verification_case_manifest.json").resolve())),
            ("求解目录", str((result_dir / "verification_runs").resolve())),
        ],
    )
    add_body(
        document,
        "限制：本模型为各向同性、纯电导率异常，不含激发极化、含水率随时间变化、地形和粗糙边界。empymod 只承担背景验证；有限三维通道由 SimPEG 与 FEniCSx 独立比较。",
        bold_lead="限制：",
    )
    add_body(
        document,
        f"最终模型指纹：{summary['model_fingerprint']}。任何配置、网格、原始数组或模型指纹缺失都会使正式 Word 构建失败。",
    )


def build_verified_report(result_dir: Path, output_path: Path) -> Path:
    result_dir = Path(result_dir)
    output_path = Path(output_path)
    summary_path = result_dir / "verification_summary.json"
    if not summary_path.is_file():
        raise VerificationGateError("formal report requires verification_summary.json")
    summary = _read_json(summary_path)
    if not summary.get("pass"):
        raise VerificationGateError(
            "formal report blocked by failed gates: "
            + ", ".join(summary.get("failed_gates", ["unknown"]))
        )
    audit = _read_json(result_dir / "model_audit.json")
    fingerprint = str(summary["model_fingerprint"])
    if audit.get("model_fingerprint") != fingerprint:
        raise VerificationGateError("model_audit.json fingerprint does not match final verification")
    for filename, _caption in FIGURES:
        if not (result_dir / filename).is_file():
            raise FileNotFoundError(result_dir / filename)

    document = Document()
    _configure_document(document, fingerprint)
    _add_cover(document, fingerprint)
    _add_model_section(document, result_dir, audit)
    _add_solver_section(document, result_dir)
    _add_gate_section(document, summary)
    _add_result_figures(document, result_dir)
    _add_reproducibility(document, result_dir, summary)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--result-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args(argv)
    print(build_verified_report(args.result_dir, args.output).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
