import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document
import pytest

from tools.fenicsx_progress_report import build_report as module


def _write_json(path: Path, value: object) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")


def _valid_inputs(tmp_path: Path) -> dict[str, Path]:
    source = tmp_path / "source.md"
    source.write_text("# 报告\n## 1. 摘要", encoding="utf-8")
    selection = tmp_path / "selection.json"
    _write_json(
        selection,
        {
            "figures": [
                {
                    "id": "figure_1",
                    "path": "figure.png",
                    "phase": "reproducibility",
                    "status": "诊断用途",
                    "placement": "appendix",
                }
            ],
            "excluded": [],
        },
    )
    inventory = tmp_path / "inventory.json"
    _write_json(
        inventory,
        {
            "images": [
                {
                    "path": "figure.png",
                    "sha256": "abc",
                    "duplicate_of": None,
                }
            ]
        },
    )
    timeline = tmp_path / "timeline.json"
    _write_json(
        timeline,
        [{"commit": "1234567890", "date": "2026-07-13", "subject": "test"}],
    )
    return {
        "source_path": source,
        "selection_path": selection,
        "inventory_path": inventory,
        "timeline_path": timeline,
        "output_path": tmp_path / "output" / "report.docx",
    }


def test_build_report_audits_reopens_and_atomically_publishes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    paths = _valid_inputs(tmp_path)
    monkeypatch.setattr(
        module,
        "audit_source",
        lambda source, selection: SimpleNamespace(ok=True, codes=[]),
    )

    def fake_build_docx(source, selection, *, project_root, output_path, timeline):
        document = Document()
        document.add_paragraph("validated")
        document.save(output_path)
        return output_path

    monkeypatch.setattr(module, "build_docx", fake_build_docx)

    summary = module.build_report(project_root=tmp_path, **paths)

    assert paths["output_path"].is_file()
    assert Document(paths["output_path"]).paragraphs[0].text == "validated"
    assert summary.figure_count == 1
    assert summary.appendix_figure_count == 1
    assert summary.git_event_count == 1
    assert len(summary.sha256) == 64
    assert not list(paths["output_path"].parent.glob("*.tmp.docx"))


def test_build_report_refuses_failed_selection_audit(tmp_path: Path) -> None:
    paths = _valid_inputs(tmp_path)
    selection = json.loads(paths["selection_path"].read_text(encoding="utf-8"))
    selection["figures"] = []
    _write_json(paths["selection_path"], selection)

    with pytest.raises(module.ReportAuditError, match="selection"):
        module.build_report(project_root=tmp_path, **paths)

    assert not paths["output_path"].exists()

