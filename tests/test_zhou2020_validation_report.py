from __future__ import annotations

import importlib.util
import hashlib
from io import BytesIO
import json
from pathlib import Path
import shutil
import zipfile
from xml.etree import ElementTree

import matplotlib
import pytest

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts/build_zhou2020_validation_report.py"
FIGURE_NAMES = (
    "fig01_model_contract.png",
    "fig02_total_fields.png",
    "fig03_reference_stability.png",
    "fig04_gate_summary.png",
    "fig05_debye_order_diagnostic.png",
)


def _load_report_script():
    spec = importlib.util.spec_from_file_location(
        "zhou_report_under_test",
        SCRIPT,
    )
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _metric(value: float, gate: float, passed: bool) -> dict[str, object]:
    return {"relative_l2": value, "gate": gate, "passed": passed}


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _canonical_digest(identity: dict[str, object]) -> str:
    encoded = json.dumps(
        identity,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _make_report_fixture(tmp_path: Path):
    run_id = "synthetic-run"
    run = (
        tmp_path
        / "generated/validation/zhou2020_grounded_wire/runs"
        / run_id
    )
    comparison = run / "comparisons/S1T1B1"
    reference = run / "reference"
    bundle = tmp_path / "publication_bundle"
    audit_dir = tmp_path / "reference_audit_hardened_v2"
    comparison.mkdir(parents=True)
    reference.mkdir(parents=True)
    bundle.mkdir()
    audit_dir.mkdir()

    total = {
        variant: {
            component: _metric(0.03, 0.05, True)
            for component in ("Ex", "Hz", "dBzdt")
        }
        for variant in ("noip", "ip")
    }
    metrics = {
        "schema": "atem3d.zhou2020.strict-comparison/v1",
        "status": "failed_with_reproducible_evidence",
        "total_field": total,
        "ip_increment": {
            "Ex": _metric(0.1632, 0.10, False),
            "Hz": _metric(0.6126, 0.10, False),
            "dBzdt": _metric(0.1140, 0.10, False),
        },
        "zero_crossings": {
            "noip": {
                "Hz": {
                    "prediction": [0.022],
                    "reference": [],
                    "passed": False,
                }
            },
            "ip": {
                "Hz": {
                    "prediction": [0.022],
                    "reference": [],
                    "passed": False,
                },
                "Ex": {
                    "prediction": [0.0576942],
                    "reference": [0.0582826],
                    "max_relative_time_error": 0.010096,
                    "passed": True,
                },
            },
        },
    }
    (comparison / "strict_comparison.json").write_text(
        json.dumps(metrics),
        encoding="utf-8",
    )
    strict_hash = _sha256(comparison / "strict_comparison.json")
    run_manifest = {
        "schema": "atem3d.zhou2020.validation-run/v1",
        "run_id": run_id,
        "comparisons": {
            "full/S1T1B1": {
                "path": "comparisons/S1T1B1/strict_comparison.json",
                "sha256": strict_hash,
                "status": "failed_with_reproducible_evidence",
            }
        },
    }
    (run / "run_manifest.json").write_text(
        json.dumps(run_manifest),
        encoding="utf-8",
    )
    input_hashes = {
        "run_manifest.json": _sha256(run / "run_manifest.json"),
        "strict_comparison.json": strict_hash,
    }
    (reference / "empymod_srcpts_convergence.json").write_text(
        json.dumps({"max_relative_difference": 0.001}),
        encoding="utf-8",
    )

    audit = {
        "schema": "atem3d.zhou2020.reference-stability/v1",
        "status": "inconclusive",
        "qwe": {
            "converged": False,
            "all_converged": False,
            "noip_total_converged": False,
            "ip_total_converged": False,
            "direct_frequency_converged": False,
        },
        "default_dlf": {"sign_changes_first20": 10},
        "stable_window": {"start_s": 5.77e-4},
        "transform_difference": {
            "default_dlf_vs_direct_qwe_relative_l2_full": 0.0644,
            "default_dlf_vs_direct_qwe_relative_l2_first20": 0.9365,
        },
        "fenicsx_vs_direct_qwe": {
            "relative_l2_full": 0.0944,
            "relative_l2_stable_window": 0.0939,
        },
    }
    diagnostic = {
        "schema": "atem3d.zhou2020.debye-order-diagnostic/v2",
        "comparison": {
            "Ex": {"debye_4_vs_16_relative_l2": 0.0920},
            "Hz": {"debye_4_vs_16_relative_l2": 0.0984},
            "dBzdt": {"debye_4_vs_16_relative_l2": 0.0640},
        },
    }
    manifest = {
        "schema": "atem3d.zhou2020.figure-bundle/v1",
        "status": "complete",
        "metadata": {
            "run_id": run_id,
            "spatial_case": "S1T1B1",
            "reference_audit_status": "inconclusive",
            "qwe_converged": False,
            "input_sha256": input_hashes,
        },
    }

    for index, name in enumerate(FIGURE_NAMES, start=1):
        fig, ax = plt.subplots(figsize=(2.0, 1.0))
        ax.plot([0.0, 1.0], [0.0, float(index)])
        ax.set_title(name)
        fig.savefig(bundle / name, dpi=80)
        plt.close(fig)

    audit_manifest = {
        "schema": "atem3d.zhou2020.reference-audit-manifest/v1",
        "status": "inconclusive",
        "input_sha256": input_hashes,
        "code_sha256": {"audit.py": "a" * 64},
        "methods": {"direct_frequency_qwe": {"ft": "qwe"}},
        "qwe": audit["qwe"],
        "artifacts": {
            "reference_stability.json": {"sha256": "b" * 64},
            "reference_stability.npz": {"sha256": "c" * 64},
        },
    }
    audit.update(
        {
            "input_sha256": input_hashes,
            "code_sha256": audit_manifest["code_sha256"],
            "methods": audit_manifest["methods"],
        }
    )
    audit_identity_core = {
        "root_relative": "output/reference_audit_hardened_v2",
        "manifest_sha256": "d" * 64,
        "manifest_schema": audit_manifest["schema"],
        "audit_schema": audit["schema"],
        "status": audit_manifest["status"],
        "artifacts": {
            name: record["sha256"]
            for name, record in audit_manifest["artifacts"].items()
        },
        "input_sha256": input_hashes,
        "code_sha256": audit_manifest["code_sha256"],
        "methods": audit_manifest["methods"],
        "qwe": audit_manifest["qwe"],
    }
    audit_identity = dict(audit_identity_core)
    audit_identity["canonical_sha256"] = _canonical_digest(audit_identity_core)
    manifest["metadata"]["reference_audit"] = audit_identity
    manifest["artifacts"] = {
        name: {
            "sha256": _sha256(bundle / name),
            "bytes": (bundle / name).stat().st_size,
        }
        for name in FIGURE_NAMES
    }

    fixture = {
        "root": tmp_path,
        "run": run,
        "bundle": bundle,
        "audit_dir": audit_dir,
        "validated_bundle": {
            "manifest": manifest,
            "diagnostic": diagnostic,
        },
        "validated_audit": {
            "manifest": audit_manifest,
            "audit": audit,
            "arrays": {},
        },
        "audit_identity": audit_identity,
    }
    return fixture


def _document_text(path: Path) -> str:
    doc = Document(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    cells = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *cells])


def _patch_validators(report, fixture, monkeypatch) -> None:
    monkeypatch.setattr(
        report,
        "_load_validated_bundle",
        lambda path, audit_path: fixture["validated_bundle"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_reference_audit",
        lambda path: fixture["validated_audit"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: fixture["audit_identity"],
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))


def test_report_consumes_validated_publication_bundle_and_audit(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    calls: list[tuple[str, Path]] = []

    def load_bundle(path: Path, audit_path: Path):
        calls.append(("bundle", Path(path)))
        assert Path(audit_path) == fixture["audit_dir"]
        return fixture["validated_bundle"]

    def load_audit(path: Path):
        calls.append(("audit", Path(path)))
        return fixture["validated_audit"]

    monkeypatch.setattr(report, "_load_validated_bundle", load_bundle)
    monkeypatch.setattr(report, "_load_validated_reference_audit", load_audit)
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: fixture["audit_identity"],
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))
    output = tmp_path / "report.docx"

    report.build_report(
        fixture["root"],
        fixture["run"],
        fixture["bundle"],
        fixture["audit_dir"],
        output,
    )

    assert calls == [
        ("bundle", fixture["bundle"]),
        ("audit", fixture["audit_dir"]),
    ]


def test_report_language_and_package_are_scientifically_fail_closed(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    monkeypatch.setattr(
        report,
        "_load_validated_bundle",
        lambda path, audit_path: fixture["validated_bundle"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_reference_audit",
        lambda path: fixture["validated_audit"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: fixture["audit_identity"],
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))
    output = tmp_path / "report.docx"

    report.build_report(
        fixture["root"],
        fixture["run"],
        fixture["bundle"],
        fixture["audit_dir"],
        output,
    )

    text = _document_text(output)
    for required in (
        "绝对值双对数",
        "绝对值不提供符号信息",
        "signed diagnostic",
        "reference-transform unstable",
        "inconclusive",
        "QWE direct 未收敛",
        "早期 DLF 震荡不是物理响应",
        "0.022 s",
        "false reversal",
        "11.40%",
        "仅为敏感性",
        "9.44%",
        "不是通过",
        "4 项相对 16 项",
        "内部敏感性",
        "failed_with_reproducible_evidence",
        "未完全通过",
        "layout rendering not performed",
        "LibreOffice unavailable",
    ):
        assert required in text, required
    for forbidden in (
        "dBz/dt 极化模块严格通过",
        "Ex、Hz 与 dBz/dt 均未同时满足 10% 严格门槛",
        "11.40% 因此正式失败",
        "9.44% 因此通过",
    ):
        assert forbidden not in text

    with zipfile.ZipFile(output) as archive:
        assert archive.testzip() is None
        names = set(archive.namelist())
        media = sorted(
            name for name in names if name.startswith("word/media/")
        )
        assert len(media) == 5
        relationships = ElementTree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        image_relationships = [
            relation
            for relation in relationships
            if relation.attrib.get("Type", "").endswith("/image")
        ]
        assert len(image_relationships) == 5
        assert all(
            relation.attrib.get("TargetMode") != "External"
            and relation.attrib["Target"].startswith("media/")
            for relation in image_relationships
        )
        assert not any(
            relation.attrib.get("TargetMode") == "External"
            for relation in relationships
        )
        document_xml = archive.read("word/document.xml")
        document_root = ElementTree.fromstring(document_xml)
        drawing_nodes = document_root.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )
        assert len(drawing_nodes) == 5


def test_report_refuses_changed_reference_audit_state(tmp_path, monkeypatch):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    changed = fixture["validated_audit"]
    changed["audit"]["status"] = "passed"
    changed["audit"]["qwe"]["converged"] = True
    monkeypatch.setattr(
        report,
        "_load_validated_bundle",
        lambda path, audit_path: fixture["validated_bundle"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_reference_audit",
        lambda path: changed,
    )
    changed_identity = dict(fixture["audit_identity"])
    changed_identity["status"] = "passed"
    changed_core = {
        key: value
        for key, value in changed_identity.items()
        if key != "canonical_sha256"
    }
    changed_identity["canonical_sha256"] = _canonical_digest(changed_core)
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: changed_identity,
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))

    try:
        report.build_report(
            fixture["root"],
            fixture["run"],
            fixture["bundle"],
            fixture["audit_dir"],
            tmp_path / "report.docx",
        )
    except ValueError as error:
        assert "re-review" in str(error)
    else:
        raise AssertionError("changed audit state must stop stale report prose")


def test_report_refuses_strict_metric_tamper_without_publishing_docx(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    _patch_validators(report, fixture, monkeypatch)
    strict = fixture["run"] / "comparisons/S1T1B1/strict_comparison.json"
    payload = json.loads(strict.read_text(encoding="utf-8"))
    payload["ip_increment"]["dBzdt"]["relative_l2"] = 0.01
    strict.write_text(json.dumps(payload), encoding="utf-8")
    output = tmp_path / "tampered.docx"

    with pytest.raises(ValueError, match="hash|identity"):
        report.build_report(
            fixture["root"],
            fixture["run"],
            fixture["bundle"],
            fixture["audit_dir"],
            output,
        )

    assert not output.exists()


def test_report_refuses_equal_looking_substituted_run_directory(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    _patch_validators(report, fixture, monkeypatch)
    substituted = tmp_path / "elsewhere" / "equal-looking-other-run"
    shutil.copytree(fixture["run"], substituted)
    output = tmp_path / "substituted.docx"

    with pytest.raises(ValueError, match="run_id|identity|path"):
        report.build_report(
            fixture["root"],
            substituted,
            fixture["bundle"],
            fixture["audit_dir"],
            output,
        )

    assert not output.exists()


def test_report_refuses_same_run_id_copied_outside_bound_run_root(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    _patch_validators(report, fixture, monkeypatch)
    substituted = tmp_path / "elsewhere" / fixture["run"].name
    shutil.copytree(fixture["run"], substituted)
    output = tmp_path / "same-id-substituted.docx"

    with pytest.raises(ValueError, match="path|identity|run root"):
        report.build_report(
            fixture["root"],
            substituted,
            fixture["bundle"],
            fixture["audit_dir"],
            output,
        )

    assert not output.exists()


def test_unbound_source_convergence_file_is_not_read(tmp_path, monkeypatch):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    _patch_validators(report, fixture, monkeypatch)
    unbound = (
        fixture["run"] / "reference/empymod_srcpts_convergence.json"
    )
    unbound.write_text("{deliberately invalid and unbound", encoding="utf-8")
    output = tmp_path / "report.docx"

    report.build_report(
        fixture["root"],
        fixture["run"],
        fixture["bundle"],
        fixture["audit_dir"],
        output,
    )

    assert output.is_file()
    assert "source quadrature difference" not in _document_text(output)


def test_report_rejects_same_inputs_but_different_audit_generation(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    different = json.loads(json.dumps(fixture["audit_identity"]))
    different["methods"]["direct_frequency_qwe"]["ft"] = "fftlog"
    core = {
        key: value
        for key, value in different.items()
        if key != "canonical_sha256"
    }
    different["canonical_sha256"] = _canonical_digest(core)
    monkeypatch.setattr(
        report,
        "_load_validated_bundle",
        lambda path, audit_path: fixture["validated_bundle"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_reference_audit",
        lambda path: fixture["validated_audit"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: different,
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))
    output = tmp_path / "mismatched-audit.docx"

    with pytest.raises(ValueError, match="audit generation identity"):
        report.build_report(
            fixture["root"],
            fixture["run"],
            fixture["bundle"],
            fixture["audit_dir"],
            output,
        )

    assert not output.exists()


def test_report_rejects_png_tampered_after_bundle_validation(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)

    def load_bundle(path, audit_path):
        (fixture["bundle"] / FIGURE_NAMES[2]).write_bytes(b"post-validation tamper")
        return fixture["validated_bundle"]

    monkeypatch.setattr(report, "_load_validated_bundle", load_bundle)
    monkeypatch.setattr(
        report,
        "_load_validated_reference_audit",
        lambda path: fixture["validated_audit"],
    )
    monkeypatch.setattr(
        report,
        "_load_validated_audit_identity",
        lambda path, validated: fixture["audit_identity"],
    )
    monkeypatch.setattr(report, "_git_state", lambda root: ("test-commit", False))
    output = tmp_path / "tampered-image.docx"

    with pytest.raises(ValueError, match="PNG|image|artifact"):
        report.build_report(
            fixture["root"],
            fixture["run"],
            fixture["bundle"],
            fixture["audit_dir"],
            output,
        )

    assert not output.exists()


def _publisher_images(fixture) -> dict[str, bytes]:
    return {
        name: (fixture["bundle"] / name).read_bytes()
        for name in FIGURE_NAMES
    }


def _minimal_docx_builder(
    images: dict[str, bytes],
    *,
    mutate_first: bool = False,
):
    def build(path: Path) -> None:
        doc = Document()
        doc.add_paragraph("atem3d.zhou2020.report/v2")
        doc.add_paragraph("failed_with_reproducible_evidence")
        for index, payload in enumerate(images.values()):
            if mutate_first and index == 0:
                payload = next(iter(images.values()))
                changed = bytearray(payload)
                changed[-1] ^= 1
                payload = bytes(changed)
            doc.add_picture(BytesIO(payload))
        doc.save(path)

    return build


def _report_debris(output: Path) -> list[Path]:
    return [
        *output.parent.glob(f".{output.name}.tmp-*"),
        *output.parent.glob(f".{output.name}.lock"),
    ]


def test_docx_publisher_rejects_lock_collision_and_preserves_other_owner(
    tmp_path,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"
    lock = report._report_lock_path(output)
    lock.write_text("other-owner", encoding="utf-8")

    with pytest.raises(FileExistsError):
        report._publish_docx(
            output,
            _minimal_docx_builder(_publisher_images(fixture)),
            _publisher_images(fixture),
        )

    assert lock.read_text(encoding="utf-8") == "other-owner"
    assert not output.exists()


def test_docx_publisher_validates_temp_before_replacing_existing_report(
    tmp_path,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"
    old = b"previous committed report"
    output.write_bytes(old)

    with pytest.raises(ValueError, match="DOCX|ZIP|package"):
        report._publish_docx(
            output,
            lambda path: path.write_bytes(b"not a docx"),
            _publisher_images(fixture),
        )

    assert output.read_bytes() == old
    assert _report_debris(output) == []


def test_docx_publisher_rejects_embedded_media_tamper(tmp_path):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"
    old = b"previous report"
    output.write_bytes(old)

    with pytest.raises(ValueError, match="media|image"):
        report._publish_docx(
            output,
            _minimal_docx_builder(
                _publisher_images(fixture),
                mutate_first=True,
            ),
            _publisher_images(fixture),
        )

    assert output.read_bytes() == old
    assert _report_debris(output) == []


def test_docx_publisher_replace_failure_preserves_old_and_cleans_debris(
    tmp_path,
    monkeypatch,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"
    old = b"previous report"
    output.write_bytes(old)
    monkeypatch.setattr(
        report.os,
        "replace",
        lambda source, target: (_ for _ in ()).throw(OSError("replace failed")),
    )

    with pytest.raises(OSError, match="replace failed"):
        report._publish_docx(
            output,
            _minimal_docx_builder(_publisher_images(fixture)),
            _publisher_images(fixture),
        )

    assert output.read_bytes() == old
    assert _report_debris(output) == []


def test_docx_publisher_keyboard_interrupt_cleans_owned_lock_and_temp(
    tmp_path,
):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"

    def interrupted(path):
        path.write_bytes(b"partial")
        raise KeyboardInterrupt()

    with pytest.raises(KeyboardInterrupt):
        report._publish_docx(
            output,
            interrupted,
            _publisher_images(fixture),
        )

    assert not output.exists()
    assert _report_debris(output) == []


def test_docx_publisher_does_not_remove_changed_lock_owner(tmp_path):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    output = tmp_path / "report.docx"
    lock = report._report_lock_path(output)

    def ownership_changed(path):
        path.write_bytes(b"partial")
        lock.write_text("new-owner", encoding="utf-8")
        raise KeyboardInterrupt()

    with pytest.raises(KeyboardInterrupt):
        report._publish_docx(
            output,
            ownership_changed,
            _publisher_images(fixture),
        )

    assert lock.read_text(encoding="utf-8") == "new-owner"
    assert not output.exists()
    lock.unlink()
    assert _report_debris(output) == []


def test_figure_paragraphs_and_captions_prevent_orphaning(tmp_path, monkeypatch):
    report = _load_report_script()
    fixture = _make_report_fixture(tmp_path)
    _patch_validators(report, fixture, monkeypatch)
    output = tmp_path / "report.docx"

    report.build_report(
        fixture["root"],
        fixture["run"],
        fixture["bundle"],
        fixture["audit_dir"],
        output,
    )

    doc = Document(output)
    captions = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Caption"
    ]
    assert len(captions) == 5
    assert all(
        paragraph.paragraph_format.keep_together
        and paragraph.paragraph_format.keep_with_next
        for paragraph in captions
    )
    picture_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph._p.xpath(".//w:drawing")
    ]
    assert len(picture_paragraphs) == 5
    assert all(
        paragraph.paragraph_format.keep_with_next
        for paragraph in picture_paragraphs
    )


def test_cli_requires_bundle_audit_and_docx_output_only():
    report = _load_report_script()
    parser = report.build_parser()

    args = parser.parse_args(
        [
            "--root",
            ".",
            "--run",
            "run",
            "--publication-bundle",
            "bundle",
            "--reference-audit",
            "audit",
            "--output",
            "report.docx",
        ]
    )
    assert args.publication_bundle == Path("bundle")
    assert args.reference_audit == Path("audit")
    assert args.output.suffix == ".docx"

    try:
        parser.parse_args(
            [
                "--root",
                ".",
                "--run",
                "run",
                "--publication-bundle",
                "bundle",
                "--reference-audit",
                "audit",
                "--output",
                "report.pdf",
            ]
        )
    except SystemExit:
        pass
    else:
        raise AssertionError("CLI must reject non-DOCX output")
