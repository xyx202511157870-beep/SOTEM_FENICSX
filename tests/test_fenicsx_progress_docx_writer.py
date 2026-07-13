from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image
import pytest

from tools.fenicsx_progress_report.docx_writer import (
    UnknownFigureError,
    build_docx,
)


def _figure(image_path: Path, *, figure_id: str = "test_figure") -> dict[str, str]:
    return {
        "id": figure_id,
        "path": str(image_path),
        "phase": "uniform_three_way",
        "status": "已验证",
        "placement": "body",
        "caption": "测试图。",
        "problem": "测试问题。",
        "change": "测试修改。",
        "result": "测试结果。",
        "source_run": "fixture",
    }


def test_build_docx_uses_compact_a4_and_embeds_figure(tmp_path: Path) -> None:
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (800, 500), "white").save(image_path)
    source = "\n".join(
        [
            "# 测试报告",
            "**报告性质：** 测试技术档案",
            "## 1. 摘要与当前状态",
            "正文包含 `dBz/dt`。",
            "[[FIGURE:test_figure]]",
        ]
    )
    output = tmp_path / "report.docx"

    build_docx(
        source,
        {"figures": [_figure(image_path)]},
        project_root=tmp_path,
        output_path=output,
    )

    document = Document(output)
    section = document.sections[0]
    assert abs(section.page_width - Mm(210)) < 1000
    assert abs(section.page_height - Mm(297)) < 1000
    assert abs(section.left_margin - Mm(17)) < 1000
    assert abs(section.right_margin - Mm(17)) < 1000
    assert len(document.inline_shapes) == 1
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "测试报告" in text
    assert "测试图" in text
    assert "FEniCSx" in section.header.paragraphs[0].text


def test_unknown_figure_directive_raises(tmp_path: Path) -> None:
    with pytest.raises(UnknownFigureError, match="missing"):
        build_docx(
            "# 报告\n## 1. 摘要与当前状态\n[[FIGURE:missing]]",
            {"figures": []},
            project_root=tmp_path,
            output_path=tmp_path / "report.docx",
        )


def test_figure_requires_caption_status_and_source(tmp_path: Path) -> None:
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (300, 200), "white").save(image_path)
    record = _figure(image_path)
    record["caption"] = ""

    with pytest.raises(ValueError, match="caption"):
        build_docx(
            "# 报告\n## 1. 摘要与当前状态\n[[FIGURE:test_figure]]",
            {"figures": [record]},
            project_root=tmp_path,
            output_path=tmp_path / "report.docx",
        )


def test_markdown_table_has_explicit_geometry_and_repeating_header(
    tmp_path: Path,
) -> None:
    source = "\n".join(
        [
            "# 报告",
            "## 1. 摘要与当前状态",
            "| 指标 | 数值 | 状态 |",
            "| --- | ---: | --- |",
            "| 中位误差 | 2.75% | 已验证 |",
        ]
    )
    output = tmp_path / "table.docx"

    build_docx(
        source,
        {"figures": []},
        project_root=tmp_path,
        output_path=output,
    )

    document = Document(output)
    table = document.tables[0]
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    assert table_width is not None
    assert int(table_width.get(qn("w:w"))) > 9000
    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    assert len(grid_cols) == 3
    assert sum(int(col.get(qn("w:w"))) for col in grid_cols) == int(
        table_width.get(qn("w:w"))
    )
    assert table.rows[0]._tr.find(qn("w:trPr")).find(qn("w:tblHeader")) is not None
    assert all(
        row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")) is not None
        for row in table.rows
    )


def test_figure_atlas_embeds_appendix_images_four_up(tmp_path: Path) -> None:
    figures = []
    for index in range(4):
        image_path = tmp_path / f"figure-{index}.png"
        Image.new("RGB", (600 + index, 400), "white").save(image_path)
        record = _figure(image_path, figure_id=f"appendix_{index}")
        record["placement"] = "appendix"
        figures.append(record)
    output = tmp_path / "atlas.docx"

    build_docx(
        "# 报告\n## 附录 B：调试效果图谱\n[[FIGURE_ATLAS]]",
        {"figures": figures},
        project_root=tmp_path,
        output_path=output,
    )

    document = Document(output)
    assert len(document.inline_shapes) == 4
    assert any(len(table.rows) == 2 and len(table.columns) == 2 for table in document.tables)


def test_major_appendix_heading_starts_on_a_new_page(tmp_path: Path) -> None:
    output = tmp_path / "appendix.docx"

    build_docx(
        "# 报告\n## 16. 当前评价\n正文。\n## 附录 B：调试效果图谱",
        {"figures": []},
        project_root=tmp_path,
        output_path=output,
    )

    document = Document(output)
    appendix = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("附录 B") and paragraph.style.name == "Heading 2"
    )
    assert appendix.paragraph_format.page_break_before is True


def test_static_contents_uses_zero_spacing_between_entries(tmp_path: Path) -> None:
    output = tmp_path / "contents.docx"

    build_docx(
        "# 报告\n## 1. 第一章\n### 1.1 小节\n正文。",
        {"figures": []},
        project_root=tmp_path,
        output_path=output,
    )

    document = Document(output)
    contents_index = next(
        index for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "目录"
    )
    entries = document.paragraphs[contents_index + 1:contents_index + 3]
    assert [paragraph.text for paragraph in entries] == ["1. 第一章", "1.1 小节"]
    assert all(paragraph.paragraph_format.space_after.pt == 0 for paragraph in entries)


def test_git_timeline_renders_commit_field(tmp_path: Path) -> None:
    output = tmp_path / "timeline.docx"

    build_docx(
        "# Report\n## Appendix A: Timeline\n[[GIT_TIMELINE]]",
        {"figures": []},
        project_root=tmp_path,
        output_path=output,
        timeline=[
            {
                "commit": "30a54d0d4b8c26273783b880a570d7740b10bb4f",
                "date": "2026-07-13T13:03:10+08:00",
                "subject": "docs: plan FEniCSx progress report",
            }
        ],
    )

    document = Document(output)
    timeline_table = document.tables[0]
    assert timeline_table.cell(1, 1).text == "30a54d0d4b"
