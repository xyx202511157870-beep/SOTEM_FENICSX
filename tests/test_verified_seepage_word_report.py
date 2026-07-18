from __future__ import annotations

import base64
import json
from pathlib import Path
import subprocess
import sys

import pytest
from docx import Document

from atem3d.seepage_verification import VerificationGateError
from tools.build_verified_seepage_word_report import (
    FIGURES,
    _gate_status_rows,
    build_verified_report,
)


def test_verified_report_cli_is_directly_executable() -> None:
    script = Path(__file__).resolve().parents[1] / "tools" / "build_verified_seepage_word_report.py"
    completed = subprocess.run(
        [sys.executable, str(script), "--help"],
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr


def test_verified_report_refuses_missing_or_failed_final_summary(
    tmp_path: Path,
) -> None:
    with pytest.raises(VerificationGateError, match="verification_summary.json"):
        build_verified_report(tmp_path, tmp_path / "report.docx")

    (tmp_path / "verification_summary.json").write_text(
        json.dumps({"pass": False, "failed_gates": ["comsol_zero_contrast"]}),
        encoding="utf-8",
    )
    with pytest.raises(VerificationGateError, match="comsol_zero_contrast"):
        build_verified_report(tmp_path, tmp_path / "report.docx")


def test_gate_rows_are_derived_from_final_json() -> None:
    rows = _gate_status_rows(
        {
            "required_gates": ["zero_contrast_simpeg", "cross_solver"],
            "gates": {
                "zero_contrast_simpeg": {
                    "available": True,
                    "pass": True,
                    "normalized_l2": [1e-10, 2e-10, 3e-10],
                },
                "cross_solver": {
                    "available": True,
                    "pass": True,
                    "median_threshold": 0.2,
                    "p95_threshold": 0.35,
                },
            },
        }
    )

    assert rows[0][0] == "zero_contrast_simpeg"
    assert rows[0][-1] == "通过"
    assert "L2" in rows[0][1]
    assert "20.0%" in rows[1][1]


def test_verified_report_contains_only_explicit_comsol_scope_exclusion(
    tmp_path: Path,
) -> None:
    fingerprint = "a" * 64
    (tmp_path / "verification_summary.json").write_text(
        json.dumps(
            {
                "pass": True,
                "failed_gates": [],
                "required_gates": [],
                "gates": {},
                "model_fingerprint": fingerprint,
            }
        ),
        encoding="utf-8",
    )
    (tmp_path / "model_audit.json").write_text(
        json.dumps(
            {
                "model_fingerprint": fingerprint,
                "channel": {
                    "size_m": [60.0, 1.0, 1.0],
                    "conductivity_s_per_m": 1.0,
                },
            }
        ),
        encoding="utf-8",
    )
    (tmp_path / "verification_case_manifest.json").write_text(
        json.dumps({"cases": []}),
        encoding="utf-8",
    )
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    for filename, _caption in FIGURES:
        (tmp_path / filename).write_bytes(png)

    output = build_verified_report(tmp_path, tmp_path / "report.docx")
    document = Document(output)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    report_text = "\n".join(text_parts)

    assert FIGURES[-1][0] == "verified_two_solver_anomaly.png"
    assert "COMSOL 不在本版正式求解、验证与报告范围内" in report_text
    assert report_text.count("COMSOL") == 1
    assert "COMSOL 三维参考" not in report_text
