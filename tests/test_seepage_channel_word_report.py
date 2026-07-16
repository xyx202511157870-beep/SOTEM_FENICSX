from __future__ import annotations

import json
from pathlib import Path

from docx import Document
import matplotlib.pyplot as plt
import pytest

from atem3d.seepage_verification import VerificationGateError


def test_report_profile_uses_model_audit_instead_of_baseline_constants() -> None:
    from tools.build_seepage_channel_word_report import _report_profile

    profile = _report_profile(
        {
            "model_fingerprint": "b" * 64,
            "channel": {
                "size_m": [60.0, 1.0, 1.0],
                "bounds_m": [[-30.0, 30.0], [-0.5, 0.5], [19.5, 20.5]],
                "theoretical_volume_m3": 60.0,
            },
            "variant": "thin_60x1x1",
        }
    )

    assert profile["size_text"] == "60 x 1 x 1 m"
    assert profile["depth_text"] == "19.5-20.5 m"
    assert profile["theoretical_volume_m3"] == 60.0
    assert profile["variant"] == "thin_60x1x1"
    assert profile["model_fingerprint"] == "b" * 64


def test_summary_conclusion_reports_failed_targets_honestly() -> None:
    from tools.build_seepage_channel_word_report import _summary_conclusion

    summary = {
        "background": {
            "SimPEG": {name: {"pass": True} for name in ("Ex", "dBzdt", "Hz")},
            "FEniCSx": {name: {"pass": True} for name in ("Ex", "dBzdt", "Hz")},
        },
        "channel_delta": {
            "Ex": {"pass": True},
            "dBzdt": {"pass": False},
            "Hz": {"pass": True},
        },
    }

    text = _summary_conclusion(summary)

    assert "dBzdt" in text
    assert "未通过" in text


def test_build_seepage_channel_report_contains_required_sections(tmp_path: Path) -> None:
    from tools.build_seepage_channel_word_report import build_report

    result_dir = tmp_path / "results"
    magnetic_dir = result_dir / "magnetic_receiver_stability"
    magnetic_dir.mkdir(parents=True)
    (result_dir / "model_audit.json").write_text(
        json.dumps(
            {
                "coordinate_convention": "z_down",
                "channel": {
                    "center_m": [0.0, 0.0, 20.0],
                    "size_m": [60.0, 1.0, 1.0],
                    "bounds_m": [[-30.0, 30.0], [-0.5, 0.5], [19.5, 20.5]],
                    "conductivity_s_per_m": 1.0,
                    "theoretical_volume_m3": 60.0,
                },
            }
        ),
        encoding="utf-8",
    )
    (magnetic_dir / "magnetic_symmetry_metrics.json").write_text(
        json.dumps(
            {
                "biot_rate": {
                    "rx3_zero_ratio": 0.004,
                    "pair_24_residual": 0.003,
                    "pair_15_residual": 0.002,
                }
            }
        ),
        encoding="utf-8",
    )
    (magnetic_dir / "magnetic_convergence_summary.json").write_text(
        json.dumps(
            {
                "passed": True,
                "selected": "biot_rate",
                "rejected": {"faraday_loop_32": ["rx3_zero_ratio"]},
            }
        ),
        encoding="utf-8",
    )
    (magnetic_dir / "mesh_symmetry_audit.json").write_text(
        json.dumps(
            {"passed": True, "solver_domain": "full", "field_mirroring": False}
        ),
        encoding="utf-8",
    )
    for stem in (
        "magnetic_receiver_comparison",
        "rx3_absolute_residual",
        "magnetic_symmetry_convergence",
    ):
        fig, axis = plt.subplots()
        axis.plot([1, 2], [1, 2])
        fig.savefig(magnetic_dir / f"{stem}.png")
        plt.close(fig)

    report_path = tmp_path / "seepage_channel_report.docx"
    built = build_report(
        result_dir=result_dir,
        output_path=report_path,
        allow_unverified_draft=True,
    )

    assert built == report_path
    document = Document(report_path)
    report_source = Path("tools/build_seepage_channel_word_report.py").read_text(
        encoding="utf-8"
    )
    assert 'add_heading(document, "4. 结果对比")' in report_source
    assert "正式报告视图：4 个接收点" in report_source
    assert "原始计算产物保留 5 点" in report_source
    assert "绝对幅值常规衰减曲线" in report_source
    assert "含通道绝对幅值常规衰减曲线" in report_source
    assert "通道异常绝对幅值常规衰减曲线" in report_source
    assert "4.3 通道带符号异常响应" not in report_source
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    table_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    text = "\n".join([*paragraph_text, *table_text])
    assert "100 m" in text
    assert "SimPEG" in text and "empymod" in text and "FEniCSx" in text
    assert "explicit_full_domain" in text
    assert "不使用单侧求解后对称镜像" in text
    assert "Rx1、Rx2、Rx4、Rx5" in text
    assert "正式报告仅展示四个非中心接收点" in text
    assert "原始五点" in text
    assert "Rx3" not in text
    assert "60 x 1 x 1 m" in text
    assert "60 x 10 x 10 m" not in text
    assert "Faraday 有限线圈" in text
    assert "候选诊断" in text
    assert "Biot-Savart Hz 时间差分（biot_rate）" in text
    assert "正式方法\nbiot_rate" in text
    assert "四面体四点 Biot-Savart" in text
    assert '("rx3_absolute_residual",' not in report_source
    assert "2 m" in text and "32" in text and "tetra4" in text
    assert len(document.tables) >= 3
    assert len(document.inline_shapes) == 1


def test_report_source_adds_channel_anomaly_diagnostic_figures() -> None:
    source = Path("tools/build_seepage_channel_word_report.py").read_text(
        encoding="utf-8"
    )
    for filename in (
        "channel_relative_anomaly.png",
        "channel_delta_signed.png",
        "channel_relative_anomaly_profiles.png",
    ):
        assert f'result_dir / "{filename}"' in source
    assert "100 x |F_channel - F_background| / |F_background|" in source
    assert 'result_dir / "channel_delta_error.png", "图 9' in source


def test_formal_report_refuses_missing_or_failed_verification_summary(
    tmp_path: Path,
) -> None:
    from tools.build_seepage_channel_word_report import build_report

    with pytest.raises(VerificationGateError, match="verification_summary.json"):
        build_report(result_dir=tmp_path, output_path=tmp_path / "report.docx")

    (tmp_path / "verification_summary.json").write_text(
        json.dumps(
            {
                "pass": False,
                "failed_gates": ["spatial_convergence"],
                "model_fingerprint": "a" * 64,
            }
        ),
        encoding="utf-8",
    )
    with pytest.raises(VerificationGateError, match="spatial_convergence"):
        build_report(result_dir=tmp_path, output_path=tmp_path / "report.docx")
